# -*- coding: utf-8 -*-
"""
Created on Thu Jan 12 05:51:13 2023

@author: SYA6COB
"""
import boto3
import threading
import logging
import os
import pymongo
from bson.json_util import dumps
import time
import shutil
from collections import defaultdict, Counter
from append_df import append_df_to_excel
from openpyxl import Workbook
from openpyxl import load_workbook
from vulnerability import vulnerability_report
# from replace_sheets import delete_sheets, replace_sheets
from sara_sheet_copy  import copy_sheets
import pandas as pd
from botocore.exceptions import ClientError
import openpyxl
from reportlab.platypus import SimpleDocTemplate

# Import the necessary modules
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import urllib
import os.path
#from test_copy_sheets import getFiles
import zipfile
import py7zr

# from Data_Base_Class_Latest import DataBase, move_list_to_df, remove_zero_rows

from create_security_doc import create_security_concepts

from Attack_trees import create_attack_vectors

from Attack_trees import create_attack_vectors_AP
#Excel to Pdf import statements


import win32com.client
import PyPDF2
from pywintypes import com_error
from reportlab.lib.pagesizes import letter
from reportlab.lib.pagesizes import A3
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from datetime import datetime
import mysql.connector


#########################################################################################################

from flask import Response

import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle,Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak
from reportlab.lib.pagesizes import landscape
from reportlab.lib.pagesizes import A4
from pdfrw import PdfReader, PdfWriter, PageMerge
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

import mimetypes
from py7zr import SevenZipFile

import fitz

from docx import Document
from docx.shared import Inches


# list to store the interfaces and session ID
job_queue = []
# tmp_queue = []
assets_list = []
methodology = ''
global dict_values

global tmp_new_sheets



wb = Workbook()

global tara_file_name
column_map = {'Assumptions': ['a', 'b'], 'MUCs': ['c', 'd'], 'DSsConsequences': ['e', 'f', 'g', 'h'], 'SecGoals': ['i', 'j', 'k'], 'ThreatDSs': ['l', 'm', 'n'], 'SecurityNeeds_AP': ['s', 't'],'ThreatEvaluation_LE':['u'],'ThreatEvaluation_AP':['v','w','x','y','z']}
columns_names = ['a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','q','r','s', 't']
starting_sheets = 0
counter_i = 0
counter_j = 0

#master_file_name = ("C:\\Users\\sya6cob\\.spyder-py3\\Master_TARA.xlsm")


# path to the TARA files
TARA_FILES_PATH = "C:\\Users\\Administrator\\Documents\\TARA_FILES\\"

# path to the TRA and RRA template
TRA_RRA_TEMPLATE_AP = "C:\\Users\\Administrator\\Documents\\sheets\\TRA_and_RRA_Template_AP.xlsm"

TRA_RRA_TEMPLATE_LE = "C:\\Users\\Administrator\\Documents\\sheets\\TRA_and_RRA_Template_LE.xlsm"

LE_Sheet_Copy = "C:\\Users\\Administrator\\Documents\\sheets\\LE_Sheet.xlsm"

AP_Sheet_Copy = "C:\\Users\\Administrator\\Documents\\sheets\\AP_Sheet.xlsm"



start_row_list = [11,32,22,14,16,57,17]
start_row_AP = [52]

# connect to the MongoDB database
# C:\Users\Administrator\Desktop\Python-Codes\global-bundle.pem
# client = pymongo.MongoClient("mongodb://m001-student:FrdDev54@sandbox-shard-00-00.hb8k2.mongodb.net:27017,sandbox-shard-00-01.hb8k2.mongodb.net:27017,sandbox-shard-00-02.hb8k2.mongodb.net:27017/?ssl=true&replicaSet=atlas-c329am-shard-0&authSource=admin&retryWrites=true/master_tara")
client = pymongo.MongoClient("mongodb://sitara:Phantom2024*@si-tara-database.cpwdcswc6fon.ap-south-1.docdb.amazonaws.com:27017/?tls=true&tlsCAFile=C%3A%5CUsers%5CAdministrator%5CDesktop%5CPython-Codes%5Cglobal-bundle.pem&replicaSet=rs0&readPreference=primary&retryWrites=false")
# "mongodb://MSECL6:Phantom2024*@master-tara.cluster-cpwdcswc6fon.ap-south-1.docdb.amazonaws.com:27017/?tls=true&tlsCAFile=D%3A%5Cmongosh-2.1.5-win32-x64%5Cmongosh-2.1.5-win32-x64%5Cbin%5Cglobal-bundle.pem&replicaSet=rs0&readPreference=primary&retryWrites=false"
database = client.get_database("master_tara")
sample_table = database.Interfaces
# print(sample_table.find())
mydb = client["master_tara"]
for coll in mydb.list_collection_names():
    print(coll)

logging.basicConfig(filename="change_stream.log", level=logging.DEBUG)

change = None

# client .master_tara.Interfaces.find()

ACCESS_KEY = 'TEST'
SECRET_KEY = 'TEST'

class DataBase:
    def __init__(self,DatabaseName):
        self._conn = mysql.connector.connect(
        host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        user="admin",
        password="phantom2022*",
         #database="tara_new"
         database=DatabaseName
        )
        
        self._cursor = self._conn.cursor()
        self.result_list = []    

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    @property
    def connection(self):
        return self._conn

    @property
    def cursor(self):
        return self._cursor

    def commit(self):
        self.connection.commit()

    def close(self, commit=True):
        if commit:
            self.commit()
        self.connection.close()

    def execute(self, sql, params=None):
        
        # print("Inside execute")
        self.cursor.execute(sql)
        
        for record in self.cursor:
            # print (record)
            self.result_list.append(record)
        # print("Ex",self.result_list)

        return self.result_list

    def fetchall(self):
        return self.cursor.fetchall()

    def fetchone(self):
        return self.cursor.fetchone()

    def query(self, sql, params=None):
        self.cursor.execute(sql, params or ())
        return self.fetchall()

def uploadDirectory(path, bucketname, session_id):
    s3 = boto3.client('s3',
                      aws_access_key_id=ACCESS_KEY,
                      aws_secret_access_key=SECRET_KEY, region_name='ap-south-1')
    for root,dirs,files in os.walk(path):
        for file in files:
            file_path = os.path.join(root, file)
            _, file_extension = os.path.splitext(file)
            # print(file_extension)
            if file == 'Attack_Trees.pdf':
                s3.upload_file(file_path,bucketname,f'Sitara_TARA_Storage/{email}/{session_id}/{file}', ExtraArgs={'ContentType': 'application/pdf'})
            elif file == 'Traceability_Graphs.pdf':
                s3.upload_file(file_path,bucketname,f'Sitara_TARA_Storage/{email}/{session_id}/{file}', ExtraArgs={'ContentType': 'application/pdf'})
            elif file_extension.lower() == '.pdf':
                s3.upload_file(file_path,bucketname,f'Sitara_TARA_Storage/{email}/{session_id}/{file}', ExtraArgs={'ContentType': 'application/pdf'})
            else:
                s3.upload_file(os.path.join(root,file),bucketname,f'Sitara_TARA_Storage/{email}/{session_id}/{file}')
            # s3.upload_file(os.path.join(root,file),bucketname,f'Sitara_TARA_Storage/{email}/{session_id}/{file}', ExtraArgs={'ContentType': 'application/pdf'})

def uploadFileToS3(folderPath, bucketName, sessionId):
    s3 = boto3.client('s3', aws_access_key_id=ACCESS_KEY,
                      aws_secret_access_key=SECRET_KEY)
    
    # seven_zip_filename = 'files.7z'
    # seven_zip_file_path = os.path.join("C:\\Users\\Administrator\\Documents\\ZIP_FILES", seven_zip_filename)

    # with SevenZipFile(seven_zip_file_path, 'w') as seven_zip_file:
    #     # Add all files in the folder to the 7z archive
    #     for root, _, files in os.walk(filePath):
    #         for file in files:
    #             file_path = os.path.join(root, file)
    #             seven_zip_file.write(file_path, os.path.relpath(file_path, filePath))

    zip_filename = 'Si-TaRA_Files.zip'

    #folder_filename = 'T'

    with zipfile.ZipFile(zip_filename, 'w') as zip_file:
        for root, dirs, files in os.walk(folderPath):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, folderPath)
                zip_file.write(file_path, arcname=arcname)

    objectKey = f'Sitara_TARA_Storage/{email}/{sessionId}/Si-TaRA_files/{zip_filename}'
    
    folderkey = f'Sitara_TARA_Storage/{email}/{sessionId}/Si-TaRA_files'

    #s3.upload_file(,bucketName,folderkey)
    s3.upload_file(zip_filename, bucketName, objectKey)


#     AWS_FILE_PATH = "C:\\Users\\Administrator\\Documents\\AWS_FILES\\"

#     os.mkdir(f'{AWS_FILE_PATH}{sessionId}')

#     s3.download_file(bucketName, objectKey, f'{AWS_FILE_PATH}{sessionId}\\TRA_and_RRA_Template_v2.3.2.xlsm')





def get_username_from_email(email):
    username = email.split('@')[0]
    username = username.replace('.', ' ')
    return username

def sendErrorEmail(ErrorMessage):
    client = boto3.client('ses',
                      aws_access_key_id=ACCESS_KEY,
                      aws_secret_access_key=SECRET_KEY, region_name='ap-south-1')
                      
    subject = "Error Occured while Generating the Tara Document"
    email = "Si.Tara@in.bosch.com"

    # Create the email message
    message = MIMEMultipart()
    message['Subject'] = subject
    message['From'] = email
    message['To'] = email
    best_regards = "\nBest Regards."
    body=MIMEText(f"Hello Admin," + '\n' + '\n' + '\t' + "The following Error has occured while generating tara document:" +'\n'+ErrorMessage+'\n'+f"{best_regards}"  )
    message.attach(body)
    # Send the email
    try:
        client.send_raw_email(
            Source=email,
            Destinations=[email],
            RawMessage={'Data': message.as_string()}
        )
        print(" Error Email sent successfully!")
    except ClientError as e:
        print(e)


def sendEmail(tara_file_path, security_concepts_docx_path,attack_trees_path,trace_path):
    mailClient = boto3.client('ses',
                      aws_access_key_id=ACCESS_KEY,
                      aws_secret_access_key=SECRET_KEY, region_name='ap-south-1')
                      
    subject = "TARA and Security Concepts Documents Has Been Generated "
    sender = "Si.Tara@in.bosch.com"
    

    # Create the email message
    message = MIMEMultipart()
    message['Subject'] = subject
    message['From'] = sender
    message['To'] = email

    # Add the email body
    assets_to_send = ",".join(assets_list)
    cloud_asset= ",".join(cloud_assets)
    interfaces_asset= ",".join(interfaces_assets)
    database_asset= ",".join(database_assets)
    application_asset= ",".join(application_assets)
    # print(assets_to_send)

    # security_concepts_docx = create_security_concepts(tara_file_name, security_concepts_docx_path)
    # print("path ios ", security_concepts_docx_path)
    email_asset_list=''
    if cloud_asset!='':
        email_asset_list=email_asset_list+"Cloud: "+cloud_asset+'\n'
    
    if interfaces_asset!='':
        email_asset_list=email_asset_list+"Interfaces: "+interfaces_asset+'\n'
    
    if database_asset!='':
        email_asset_list=email_asset_list+"Database: "+ database_asset +'\n'
    
    if application_asset!='':
        email_asset_list=email_asset_list+"Application: "+application_asset+'\n'

    vulnerability_report_path=f"{TARA_FILES_PATH}{session_id}"+"\\vulnerability_report.xlsx"
    vulnerability_report_not_available=''

    if(vul_report_gen==True and os.path.exists(vulnerability_report_path)==False):
        vulnerability_report_not_available='\nNo vulnerability was found for the selected assests\n'


    name = get_username_from_email(email)
    best_regards = "\nBest Regards,"
    aigenerated="\nSome of the context are generated using AI, Please review before using it."
    body = MIMEText(f"Hello {name.upper()}," + '\n' + '\n' + '\t' + "TaRA and Security Concepts documents are generated for the following Assets using "+methodology+" methodology : \n" +'\n'+email_asset_list+vulnerability_report_not_available+aigenerated+'\n\n' + f"{best_regards}" + '\n' + "Si-TARA Team.")
    # print(body)
    message.attach(body)

    # Add the attachment
    # attachment_path = "C:\\Users\\Administrator\\Desktop\\Python-Codes\\workiong-code.txt"
    tara_attachment_name = os.path.basename(tara_file_path)
    
    if(tara_attachment_name == "TRA_and_RRA_Template_AP.xlsm"):
        tara_attachment_name = "Si-TaRA_Threat_and_Risk_Analysis_AP.xlsm"
    else:
        tara_attachment_name = "Si-TaRA_Threat_and_Risk_Analysis_LE.xlsm"

    # print("security_concepts_docx path is", security_concepts_docx)
    #security_concepts_attachment_name = os.path.basename(security_concepts_docx_path)

    # print(security_concepts_attachment_name)'

    #pdf_loc = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Final_PDF_Report.pdf'

    #PDF_path = os.path.basename(final_pdf_report)

    # print("new variable path is",attack_trees_path)
    ATTACK_TREES_path = os.path.basename(attack_trees_path)
    #ATTACK_TREES_path = attack_trees_path 
    # print("final attack path is",ATTACK_TREES_path)
        
    TRACEBILITY_path = os.path.basename(trace_path)


    with open(tara_file_path, 'rb') as tara_attachment:
        tara_attachment_content = tara_attachment.read()

    with open(security_concepts_docx_path, 'rb') as security_attachment:
        security_attachment_content = security_attachment.read()

    # print(final_pdf_report)

    #final_report = shutil.copy(final_pdf_report, f"{TARA_FILES_PATH}{session_id}")

    # with open(final_pdf_report, 'rb') as PDF_attachment:
    #     pdf_attachment_content = PDF_attachment.read()

    with open(attack_trees_path, 'rb') as attacktrees_attachment:
        attacktrees_attachment_content = attacktrees_attachment.read()

    attack_tree_final = shutil.copy(attack_trees_path, f"{TARA_FILES_PATH}{session_id}")

    tara_attachment = MIMEApplication(tara_attachment_content, Name=tara_attachment_name)
    security_attachment = MIMEApplication(security_attachment_content, Name='Si-TaRA_Security_Concepts.docx')
    #PDF_attachment = MIMEApplication(pdf_attachment_content, Name=PDF_path)
    
    attacktrees_attachment = MIMEApplication(attacktrees_attachment_content, Name=ATTACK_TREES_path)

    # print(change["fullDocument"]["docType"])
    


    ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



    directory_path3 = f"{ZIPFILES_PATH}{session_id}"

    if not os.path.exists(directory_path3):
        os.mkdir(directory_path3)
        print(f"Directory '{directory_path3}' created successfully.")
    else:
        print(f"Directory '{directory_path3}' already exists.")   




    #os.mkdir(f"{ZIPFILES_PATH}{session_id}")


    
    ZIPFOLDER_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FOLDERS\\"


    directory_path4 = f"{ZIPFOLDER_PATH}{session_id}"

    if not os.path.exists(directory_path4):
        os.mkdir(directory_path4)
        print(f"Directory '{directory_path4}' created successfully.")
    else:
        print(f"Directory '{directory_path4}' already exists.")   


    #os.mkdir(f"{ZIPFOLDER_PATH}{session_id}")
    # Usage example
    folder_to_zip = f"C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\"
    #folder_to_zip = f"{TARA_FILES_PATH}{session_id}"
    #output_zip_file = f'C:\\Users\\Administrator\\Documents\\ZIP_FOLDERS\\{session_id}\\Si-TaRA_DOCS.zip'  # Using a raw string for the file path
    
    output_zip_file = f'C:\\Users\\Administrator\\Documents\\ZIP_FOLDERS\\{session_id}\\Attack_Trees.7z'

    output_zip_trace_file = f'C:\\Users\\Administrator\\Documents\\ZIP_FOLDERS\\{session_id}\\Tracebility.7z'

    #zip_folder(folder_to_zip, output_zip_file)
    #compress_directory(folder_to_zip, output_zip_file)    # before Attack Trees without header and footer


    compress_file(attack_trees_path,output_zip_file)    

    compress_file(trace_path,output_zip_trace_file)

    #compress_directory(output_zip_file)


    zip_path = "Si-TaRA_Attack_Trees_PDF.7z"

    zip_path1 = "Si-TaRA_Tracebility_Graphs_PDF.7z"

    #zip_path = os.path.basename(output_zip_file)

    with open(output_zip_file, 'rb') as ZIP_attachment:
        zip_attachment_content = ZIP_attachment.read()

    ZIP_attachment = MIMEApplication(zip_attachment_content, Name=zip_path)

    
    with open(output_zip_trace_file, 'rb') as ZIP_Trace_attachment:
        zip_trace_attachment_content = ZIP_Trace_attachment.read()

    ZIP_Trace_attachment = MIMEApplication(zip_trace_attachment_content, Name=zip_path1)


    for key in change["fullDocument"]["docType"]:
        if key == 'AP .xlsm' or key == 'LE .xlsm': #.xlsm 
            message.attach(tara_attachment)
            
            if key == 'AP .xlsm':
                shutil.copy(tara_file_path,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.xlsm")
                AP_COUNT = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_xl": AP_COUNT["AP_xl"] + 1}})
            else:
                shutil.copy(tara_file_path,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_LE.xlsm")
                LE_COUNT = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"LE_xl": LE_COUNT["LE_xl"] + 1}})


        if key == 'AP SCD' or key == 'LE SCD': # .xlsm scd
            message.attach(security_attachment)
            shutil.copy(security_concepts_docx_path,f"{ZIPFILES_PATH}{session_id}")
            if key == 'AP SCD':
                # print("insude SCD")
                AP_SCD = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_SCD": AP_SCD["AP_SCD"] + 1}})
            else:
                LE_SCD = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"LE_SCD": LE_SCD["LE_SCD"] + 1}})



        


        if key == 'AP .PDF' or key == 'LE .PDF': # .pdf
          #  message.attach(PDF_attachment)
            #shutil.copy(final_pdf_report,f"{ZIPFILES_PATH}{session_id}")
            if key == 'AP .PDF':
                # print('inside PDF')

                pdf_data = client.tara_pdf.PDF.find_one({'filename': session_id})

                if pdf_data:

                # Get the PDF content        
                    pdf_content = pdf_data['pdf']


                    # print(pdf_content,"This is pdf_content")

                #file = Response(pdf_content, content_type='application/pdf')


                    path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf' 


                    try:
                        with open(path,'wb') as f:
                            f.write(pdf_content)
                    except Exception as e:
                        print(e)


                    newpath = header_footer4(path)

                
                # # Serve the PDF as a response with Content-Type set to application/pdf
                
                    path1 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP1.pdf'

                    finalpath = merge_pdfs_from_folder2(attack_trees_path,trace_path,path1)



                    shutil.copy(finalpath,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf")
                    # client.tara_pdf.PDF.find_one_and_update({})
                   # shutil.copy(output_zip_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Attack_Trees.7z")
                   # shutil.copy(output_zip_trace_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Traceability_Graphs.7z")
                

                #shutil.copy(final_pdf_report,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf")
                
                
                AP_PDF = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_PDF": AP_PDF["AP_PDF"] + 1}})
            else:
             #   shutil.copy(final_pdf_report,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_LE.pdf")
                LE_PDF = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"LE_PDF": LE_PDF["LE_PDF"] + 1}})

        if key == 'AP AT':
            # print("inside at")
            message.attach(ZIP_attachment)
            shutil.copy(output_zip_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Attack_Trees.7z")
            AP_AT = client.master_tara.UserData.find_one({"email": email})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_AT": AP_AT["AP_AT"] + 1}})

        if key == 'AP TR' or key == 'LE TR':
            message.attach(ZIP_Trace_attachment)

            if key =='AP TR':
                shutil.copy(output_zip_trace_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Traceability_Graphs.7z")
                AP_TR = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email":email},{"$set": {"AP_TR": AP_TR["AP_TR"] + 1}})
            else:
                shutil.copy(output_zip_trace_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Traceability_Graphs.7z")
                LE_TR = client.master_tara.UserData.find_one({"email": email})
                client.master_tara.UserData.find_one_and_update({"email":email},{"$set": {"LE_TR": LE_TR["LE_TR"] + 1}})

    if(vul_report_gen==True and os.path.exists(vulnerability_report_path)):
        with open(vulnerability_report_path, 'rb') as attacktrees_attachment:
            vulnerability_attachment_content = attacktrees_attachment.read()
        vulnerability_attachment=MIMEApplication(vulnerability_attachment_content, Name='vulnerability_document.xlsx')
        message.attach(vulnerability_attachment)
        shutil.copy(vulnerability_report_path,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Vulnerability_Report.xlsx")
        # print('vulnerability report attached')
        if methodology == "Attack Potential":
            AP_VD = client.master_tara.UserData.find_one({"email": email})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_VD": AP_VD["AP_VD"] + 1}})
        else:
            LE_VD = client.master_tara.UserData.find_one({"email": email})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"LE_VD": LE_VD["LE_VD"] + 1}})

    # print(email)


    
    uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
    uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)



    # # Send the email
    # try:
    #     mailClient.send_raw_email(
    #         Source=sender,
    #         Destinations=[email],
    #         RawMessage={'Data': message.as_string()}
    #     )
    #     print("Email sent successfully!")
    # except ClientError as e:
    #     print(e)


##########################################################################################################################

def header_footer_AP_LE(pdf_path):
    # Load your PDF
    #pdf_path = 'Si-TaRA_Threat_and_Risk_Analysis_AP_new.pdf'
    pdf_document = fitz.open(pdf_path)
    
    # Load your images
    header = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
    footer = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'
    header_image = fitz.Pixmap(header)
    footer_image = fitz.Pixmap(footer)
    
    # Define the height for the header and footer
    header_height = 70  # Adjust this value as needed
    footer_height = 30  # Adjust this value as needed
    
    new_pdf_document = fitz.open()
 
    for page_num in range(len(pdf_document)):
        original_page = pdf_document[page_num]
    
        if page_num == 0:  # For the first page
            new_page_height = original_page.rect.height
        else:  # For other pages
            new_page_height = original_page.rect.height + header_height + footer_height
    
        new_page = new_pdf_document.new_page(width=original_page.rect.width, height=new_page_height)
    
        if page_num != 0:  # Skip header and footer for the first page
            header_rect = fitz.Rect(0, 0, original_page.rect.width, header_height)
            footer_rect = fitz.Rect(0, new_page_height - footer_height, original_page.rect.width, new_page_height)
    
            new_page.insert_image(header_rect, pixmap=header_image, keep_proportion=False)
            new_page.insert_image(footer_rect, pixmap=footer_image, keep_proportion=False)
    
        # Increase the zoom factor for better clarity
        matrix = fitz.Matrix(1.5, 1.5)  # A value like 1.5 or 2 can be used
        if page_num == 0:
            pixmap = original_page.get_pixmap(matrix=matrix)
            new_page.insert_image((0, 0, original_page.rect.width, original_page.rect.height), pixmap=pixmap)
        else:
            pixmap = original_page.get_pixmap(matrix=matrix)
            new_page.insert_image((0, header_height, original_page.rect.width, original_page.rect.height + header_height), pixmap=pixmap)
    
    new_pdf_path = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\ribbon_added.pdf'
    new_pdf_document.save(new_pdf_path, garbage=4, deflate=True)  # Clean up and compress the file
    pdf_document.close()
    new_pdf_document.close()
    return new_pdf_path







    





def header_footer(pdf_pathloc): 
    # Load the original PDF
    pdf_path = pdf_pathloc
    #pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\Si-TaRA_Threat_and_Risk_Analysis_AP_new.pdf'
    original_pdf = fitz.open(pdf_path)
 
    # Create a new PDF document for the modified version
    modified_pdf = fitz.open()
 
    # Load your header and footer images
    header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
    footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'
    header_image = fitz.Pixmap(header_image_path)
    footer_image = fitz.Pixmap(footer_image_path)
 
    # Calculate the scale for the images to fit the page width and maintain aspect ratio
    page_width = original_pdf[0].rect.width  # Assuming all pages have the same width
    header_scale_factor = page_width / header_image.width
    footer_scale_factor = page_width / footer_image.width
    header_height = header_image.height * header_scale_factor
    footer_height = footer_image.height * footer_scale_factor
 
    # Process each page
    for page_num in range(len(original_pdf)):
        old_page = original_pdf[page_num]
 
        # Calculate new page height
        new_page_height = old_page.rect.height + (header_height if page_num >= 1 else 0)
 
        # Create a new page in the modified document with adjusted height
        new_page = modified_pdf.new_page(-1, width=old_page.rect.width, height=new_page_height)
 
        # Define the area where the old page content will be placed
        content_rect = fitz.Rect(0, header_height if page_num >= 1 else 0, page_width, new_page_height)
 
        # Copy the content from the old document to the new page
        new_page.show_pdf_page(content_rect, original_pdf, page_num)
 
        # Add header and footer if not the first page
        if page_num >= 1:
            new_page.insert_image(fitz.Rect(0, 0, page_width, header_height), pixmap=header_image)
            footer_y_position = new_page.rect.height - footer_height
            new_page.insert_image(fitz.Rect(0, footer_y_position, page_width, new_page.rect.height), pixmap=footer_image)
 
    # Save the new PDF with headers and footers
    new_pdf_path = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\ribbon_added.pdf'        
    #new_pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\modified_pdf_with_header_footer_new1.pdf'
    modified_pdf.save(new_pdf_path)
    modified_pdf.close()
    original_pdf.close()
    return new_pdf_path



def header_footer1(pdf_pathloc1): 
    # Load the original PDF
    pdf_path1 = pdf_pathloc1
    #pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\Si-TaRA_Threat_and_Risk_Analysis_AP_new.pdf'
    original_pdf = fitz.open(pdf_path1)
 
    # Create a new PDF document for the modified version
    modified_pdf1 = fitz.open()
 
    # Load your header and footer images
    header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
    footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'
    header_image = fitz.Pixmap(header_image_path)
    footer_image = fitz.Pixmap(footer_image_path)
 
    # Calculate the scale for the images to fit the page width and maintain aspect ratio
    page_width = original_pdf[0].rect.width  # Assuming all pages have the same width
    header_scale_factor = page_width / header_image.width
    footer_scale_factor = page_width / footer_image.width
    header_height = header_image.height * header_scale_factor
    footer_height = footer_image.height * footer_scale_factor
 
    # Process each page
    for page_num in range(len(original_pdf)):
        old_page = original_pdf[page_num]
 
        # Calculate new page height
        new_page_height = old_page.rect.height + (header_height if page_num >= 1 else 0)
 
        # Create a new page in the modified document with adjusted height
        new_page = modified_pdf1.new_page(-1, width=old_page.rect.width, height=new_page_height)
 
        # Define the area where the old page content will be placed
        content_rect = fitz.Rect(0, header_height if page_num >= 1 else 0, page_width, new_page_height)
 
        # Copy the content from the old document to the new page
        new_page.show_pdf_page(content_rect, original_pdf, page_num)
 
        # Add header and footer if not the first page
        if page_num >= 1:
            new_page.insert_image(fitz.Rect(0, 0, page_width, header_height), pixmap=header_image)
            footer_y_position = new_page.rect.height - footer_height
            #new_page.insert_image(fitz.Rect(0, footer_y_position, page_width, new_page.rect.height), pixmap=footer_image)
 
    # Save the new PDF with headers and footers
    new_pdf_path1 = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Attack_Trees.pdf'   

    #new_pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\modified_pdf_with_header_footer_new1.pdf'
    modified_pdf1.save(new_pdf_path1)
    modified_pdf1.close()
    original_pdf.close()

    return new_pdf_path1


def header_footer2(pdf_pathloc2): 
    # Load the original PDF
    pdf_path2 = pdf_pathloc2
    #pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\Si-TaRA_Threat_and_Risk_Analysis_AP_new.pdf'
    original_pdf = fitz.open(pdf_path2)
 
    # Create a new PDF document for the modified version
    modified_pdf2 = fitz.open()
 
    # Load your header and footer images
    header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
    footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'
    header_image = fitz.Pixmap(header_image_path)
    footer_image = fitz.Pixmap(footer_image_path)
 
    # Calculate the scale for the images to fit the page width and maintain aspect ratio
    page_width = original_pdf[0].rect.width  # Assuming all pages have the same width
    header_scale_factor = page_width / header_image.width
    footer_scale_factor = page_width / footer_image.width
    header_height = header_image.height * header_scale_factor
    footer_height = footer_image.height * footer_scale_factor
 
    # Process each page
    for page_num in range(len(original_pdf)):
        old_page = original_pdf[page_num]
 
        # Calculate new page height
        new_page_height = old_page.rect.height + (header_height if page_num >= 1 else 0) + (footer_height if page_num >= 1 else 0)
 
        # Create a new page in the modified document with adjusted height
        new_page = modified_pdf2.new_page(-1, width=old_page.rect.width, height=new_page_height)
 
        # Define the area where the old page content will be placed
        content_rect = fitz.Rect(0, header_height if page_num >= 1 else 0, page_width, new_page_height)
 
        # Copy the content from the old document to the new page
        new_page.show_pdf_page(content_rect, original_pdf, page_num)
 
        # Add header and footer if not the first page
        if page_num >= 1:
            new_page.insert_image(fitz.Rect(0, 0, page_width, header_height), pixmap=header_image)
            footer_y_position = new_page.rect.height - footer_height
            new_page.insert_image(fitz.Rect(0, footer_y_position, page_width, new_page.rect.height), pixmap=footer_image)
 
    # Save the new PDF with headers and footers
    new_pdf_path2 = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Traceability_Graphs.pdf'        
    #new_pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\modified_pdf_with_header_footer_new1.pdf'
    modified_pdf2.save(new_pdf_path2)
    modified_pdf2.close()
    original_pdf.close()
    return new_pdf_path2





def header_footer4(pdf_pathloc4): 
    # Load the original PDF
    pdf_path4 = pdf_pathloc4
    #pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\Si-TaRA_Threat_and_Risk_Analysis_AP_new.pdf'
    original_pdf = fitz.open(pdf_path4)
 
    # Create a new PDF document for the modified version
    modified_pdf4 = fitz.open()
 
    # Load your header and footer images
    header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
    footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'
    header_image = fitz.Pixmap(header_image_path)
    footer_image = fitz.Pixmap(footer_image_path)
 
    # Calculate the scale for the images to fit the page width and maintain aspect ratio
    page_width = original_pdf[0].rect.width  # Assuming all pages have the same width
    header_scale_factor = page_width / header_image.width
    footer_scale_factor = page_width / footer_image.width
    header_height = header_image.height * header_scale_factor
    footer_height = footer_image.height * footer_scale_factor
 
    # Process each page
    for page_num in range(len(original_pdf)):
        old_page = original_pdf[page_num]
 
        # Calculate new page height
        new_page_height = old_page.rect.height + (header_height if page_num >= 1 else 0) + (footer_height if page_num >= 1 else 0)
 
        # Create a new page in the modified document with adjusted height
        new_page = modified_pdf4.new_page(-1, width=old_page.rect.width, height=new_page_height)
 
        # Define the area where the old page content will be placed
        content_rect = fitz.Rect(0, header_height if page_num >= 1 else 0, page_width, new_page_height)
 
        # Copy the content from the old document to the new page
        new_page.show_pdf_page(content_rect, original_pdf, page_num)
 
        # Add header and footer if not the first page
        if page_num >= 1:
            new_page.insert_image(fitz.Rect(0, 0, page_width, header_height), pixmap=header_image)
            footer_y_position = new_page.rect.height - footer_height
            new_page.insert_image(fitz.Rect(0, footer_y_position, page_width, new_page.rect.height), pixmap=footer_image)
 
    # Save the new PDF with headers and footers
    new_pdf_path4 = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf'        
    #new_pdf_path = 'C:\\Users\\UGY2KOR\\Desktop\\modified_pdf_with_header_footer_new1.pdf'
    modified_pdf4.save(new_pdf_path4)
    modified_pdf4.close()
    original_pdf.close()
    return new_pdf_path4



# def header_footer3(doc_path):
# # Path to your Word document
#         #doc_path = 'test.docx'
#         #header_image_path = 'header.png'
#         #footer_image_path = 'footer.png'
#         header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
#         footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'

#         # Load the document
#         document = Document(doc_path)
        
#         # Add header and footer to each section of the document
#         for section in document.sections:
#             # Add header
#             header = section.header
#             paragraph = header.paragraphs[0]
#             run = paragraph.add_run()
#             run.add_picture(header_image_path, width=Inches(6))  # Adjust width as needed
        
#             # Add footer
#             footer = section.footer
#             paragraph = footer.paragraphs[0]
#             run = paragraph.add_run()
#             run.add_picture(footer_image_path, width=Inches(6))  # Adjust width as needed
        
#         # Save the modified document
#         new_doc_path3 = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Security_Concepts.docx'
#         document.save(new_doc_path3)
#         return new_doc_path3



def header_footer3(doc_path):
# Path to your Word document
        #doc_path = 'test.docx'
        #header_image_path = 'header.png'
        #footer_image_path = 'footer.png'
        header_image_path = 'C:\\Users\\Administrator\\Desktop\\header_sitara.png'
        footer_image_path = 'C:\\Users\\Administrator\\Desktop\\footer_sitara.png'

        # Load the document
        document = Document(doc_path)
        
        # Add header and footer to each section of the document
        for section in document.sections:
            # Add header
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.add_run().add_break()
        
        
            run = paragraph.add_run()
            run.add_picture(header_image_path, width=Inches(7.08),height=Inches(0.5))  # Adjust width as needed
        
            # Add footer
            footer = section.footer
            paragraph = footer.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(footer_image_path, width=Inches(6))  # Adjust width as needed
        
        # Save the modified document
        new_doc_path3 = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Security_Concepts.docx'
        document.save(new_doc_path3)
        return new_doc_path3















################################################################################################################################
def compress_directory(dir_path, output_path):
    with py7zr.SevenZipFile(output_path, 'w') as archive:
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, dir_path)
                archive.write(file_path, arcname=arcname)
 

def compress_file(input_file, output_path):
    with py7zr.SevenZipFile(output_path, 'w') as archive:
        arcname = os.path.basename(input_file)
        archive.write(input_file, arcname=arcname)


#######################################################################################################################


def zip_folder(folder_path, output_zip):
    """
    Zips the contents of the specified folder.
 
    Args:
    folder_path (str): The path of the folder to be zipped.
    output_zip (str): The path of the output zip file, including .zip extension.
 
    Returns:
    None
    """
    try:
        with zipfile.ZipFile(output_zip, 'w', zipfile.ZIP_LZMA) as zipf:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    zipf.write(file_path, os.path.relpath(file_path, folder_path))
        #print(f"Folder '{folder_path}' has been zipped to '{output_zip}' successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")
 


#################################################################################################################

def merge_multiple_pdfs(pdf_paths, output_path):
    pdf_writer = PyPDF2.PdfWriter()

    for pdf_path in pdf_paths:
        pdf = PyPDF2.PdfReader(pdf_path)
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            pdf_writer.add_page(page)

    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

def create_title_page(title, subtitle,man_name,projectname,taratype, output_path,image_path):
    c = canvas.Canvas(output_path, pagesize=landscape(letter))

    current_date = change["fullDocument"]["dateTime"]
    #background_color = HexColor("#330033")  # Dark grayish-blue background
    #text_color = HexColor("#00FFFF")

    background_color = HexColor("#FFFFFF")  # Dark grayish-blue background   #output 33
    text_color = HexColor("#330033")      #output 33

    #background_color = HexColor("#ADD8E6")  # Dark grayish-blue background   #output 34
    #text_color = HexColor("#000000")           #output 34

    #font_path = "C:\\Users\\Administrator\\Desktop\\LeagueSpartan-Bold.ttf"  # Replace this with your .ttf font file path


    #pdfmetrics.registerFont(TTFont('CustomFont', font_path))
    

    c.setFillColor(background_color)
    c.rect(0, 0, landscape(letter)[0], landscape(letter)[1], fill=1)




    # Title
    c.setFillColor(text_color)

    #c.drawImage(image_path, 400, 600, width=30, height=50)

    c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-130, width=800, height=350)
 
    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title
 
    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle
 
    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
   
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
 
 
    




    #c.drawImage(image_path, landscape(letter)[0]/2+60, landscape(letter)[0]/2, width=50, height=70)

    #c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-100, width=800, height=800) 

    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
    
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
    title_font_size = 31 if len(projectname) < 20 else 25  # Adjust as needed
    
    manager_font_size = 10 if len(man_name) < 20 else 8


    #c.setFont("Helvetica-Bold", 31)
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
   
    #c.setFont("Helvetica", 10)
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, man_name)








    # c.setFont("Helvetica-Bold", 43)
    # c.drawCentredString(landscape(letter)[0]/2-20, landscape(letter)[0]/2, title)  # Position the title

    # # Subtitle
    # c.setFont("Helvetica", 13)
    # c.drawCentredString(landscape(letter)[0]/2+40, landscape(letter)[0]/2-30, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    
    c.setFont("Helvetica", 10)   
    c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-180, taratype) 

    c.setFont("Helvetica", 10)
    c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date 

    c.setFont("Helvetica", manager_font_size)
    c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, "Security Manager Name: "+man_name)  #Position the Security Manager Name


    c.save()



def create_title_page_AP(man_name,projectname, output_path,image_path):
    c = canvas.Canvas(output_path, pagesize=landscape(letter))

    current_date = change["fullDocument"]["dateTime"]
    #background_color = HexColor("#330033")  # Dark grayish-blue background
    #text_color = HexColor("#00FFFF")

    background_color = HexColor("#FFFFFF")  # Dark grayish-blue background   #output 33
    text_color = HexColor("#330033")      #output 33

    #background_color = HexColor("#ADD8E6")  # Dark grayish-blue background   #output 34
    #text_color = HexColor("#000000")           #output 34

    #font_path = "C:\\Users\\Administrator\\Desktop\\LeagueSpartan-Bold.ttf"  # Replace this with your .ttf font file path


    #pdfmetrics.registerFont(TTFont('CustomFont', font_path))
    

    c.setFillColor(background_color)
    c.rect(0, 0, landscape(letter)[0], landscape(letter)[1], fill=1)




    # Title
    c.setFillColor(text_color)

    #c.drawImage(image_path, 400, 600, width=30, height=50)

    c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-130, width=800, height=350)
 
    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title
 
    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle
 
    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
   
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
 
 
    




    #c.drawImage(image_path, landscape(letter)[0]/2+60, landscape(letter)[0]/2, width=50, height=70)

    #c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-100, width=800, height=800) 

    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
    
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
    title_font_size = 31 if len(projectname) < 20 else 25  # Adjust as needed
    
    manager_font_size = 10 if len(man_name) < 20 else 8


    #c.setFont("Helvetica-Bold", 31)
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
   
    #c.setFont("Helvetica", 10)
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, man_name)








    # c.setFont("Helvetica-Bold", 43)
    # c.drawCentredString(landscape(letter)[0]/2-20, landscape(letter)[0]/2, title)  # Position the title

    # # Subtitle
    # c.setFont("Helvetica", 13)
    # c.drawCentredString(landscape(letter)[0]/2+40, landscape(letter)[0]/2-30, subtitle)  # Position the subtitle

    attacktrees = "Attack Trees"
    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-230, attacktrees)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    


    # c.setFont("Helvetica", 10)   
    # c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-180, taratype) 

    c.setFont("Helvetica", 10)
    c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date 

    c.setFont("Helvetica", manager_font_size)
    c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, "Security Manager Name: "+man_name)  #Position the Security Manager Name


    c.save()



def create_title_page_TR(man_name,projectname, output_path,image_path):
    c = canvas.Canvas(output_path, pagesize=landscape(letter))

    current_date = change["fullDocument"]["dateTime"]
    #background_color = HexColor("#330033")  # Dark grayish-blue background
    #text_color = HexColor("#00FFFF")

    background_color = HexColor("#FFFFFF")  # Dark grayish-blue background   #output 33
    text_color = HexColor("#330033")      #output 33

    #background_color = HexColor("#ADD8E6")  # Dark grayish-blue background   #output 34
    #text_color = HexColor("#000000")           #output 34

    #font_path = "C:\\Users\\Administrator\\Desktop\\LeagueSpartan-Bold.ttf"  # Replace this with your .ttf font file path


    #pdfmetrics.registerFont(TTFont('CustomFont', font_path))
    

    c.setFillColor(background_color)
    c.rect(0, 0, landscape(letter)[0], landscape(letter)[1], fill=1)




    # Title
    c.setFillColor(text_color)

    #c.drawImage(image_path, 400, 600, width=30, height=50)

    c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-130, width=800, height=350)
 
    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title
 
    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle
 
    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
   
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
 
 
    




    #c.drawImage(image_path, landscape(letter)[0]/2+60, landscape(letter)[0]/2, width=50, height=70)

    #c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-100, width=800, height=800) 

    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
    
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
    title_font_size = 31 if len(projectname) < 20 else 25  # Adjust as needed
    
    manager_font_size = 10 if len(man_name) < 20 else 8


    #c.setFont("Helvetica-Bold", 31)
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
   
    #c.setFont("Helvetica", 10)
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, man_name)








    # c.setFont("Helvetica-Bold", 43)
    # c.drawCentredString(landscape(letter)[0]/2-20, landscape(letter)[0]/2, title)  # Position the title

    # # Subtitle
    # c.setFont("Helvetica", 13)
    # c.drawCentredString(landscape(letter)[0]/2+40, landscape(letter)[0]/2-30, subtitle)  # Position the subtitle

    
    attacktrees = "Tracebility Graphs"

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-230, attacktrees)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    


    # c.setFont("Helvetica", 10)   
    # c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-180, taratype) 

    c.setFont("Helvetica", 10)
    c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date 

    c.setFont("Helvetica", manager_font_size)
    c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, "Security Manager Name: "+man_name)  #Position the Security Manager Name


    c.save()




def create_title_page_TaRA(man_name,projectname, output_path,image_path):
    c = canvas.Canvas(output_path, pagesize=landscape(A4))

    current_date = change["fullDocument"]["dateTime"]
    #background_color = HexColor("#330033")  # Dark grayish-blue background
    #text_color = HexColor("#00FFFF")

    background_color = HexColor("#FFFFFF")  # Dark grayish-blue background   #output 33
    text_color = HexColor("#330033")      #output 33

    #background_color = HexColor("#ADD8E6")  # Dark grayish-blue background   #output 34
    #text_color = HexColor("#000000")           #output 34

    #font_path = "C:\\Users\\Administrator\\Desktop\\LeagueSpartan-Bold.ttf"  # Replace this with your .ttf font file path


    #pdfmetrics.registerFont(TTFont('CustomFont', font_path))
    

    c.setFillColor(background_color)
    c.rect(0, 0, landscape(A4)[0], landscape(A4)[1], fill=1)




    # Title
    c.setFillColor(text_color)

    #c.drawImage(image_path, 400, 600, width=30, height=50)

    c.drawImage(image_path, landscape(A4)[0]/2-400, landscape(A4)[0]/2-130, width=800, height=350)
 
    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title
 
    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle
 
    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
   
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
 
 
    




    #c.drawImage(image_path, landscape(letter)[0]/2+60, landscape(letter)[0]/2, width=50, height=70)

    #c.drawImage(image_path, landscape(letter)[0]/2-400, landscape(letter)[0]/2-100, width=800, height=800) 

    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
    
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """
    title_font_size = 31 if len(projectname) < 20 else 25  # Adjust as needed
    
    manager_font_size = 10 if len(man_name) < 20 else 8


    #c.setFont("Helvetica-Bold", 31)
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-180, projectname)  # Position the title
   
    #c.setFont("Helvetica", 10)
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2+300, landscape(letter)[0]/2-360, current_date)  #Position the date
 
    # c.setFont("Helvetica", 10)
    # c.drawCentredString(landscape(letter)[0]/2-240, landscape(letter)[0]/2-360, man_name)








    # c.setFont("Helvetica-Bold", 43)
    # c.drawCentredString(landscape(letter)[0]/2-20, landscape(letter)[0]/2, title)  # Position the title

    # # Subtitle
    # c.setFont("Helvetica", 13)
    # c.drawCentredString(landscape(letter)[0]/2+40, landscape(letter)[0]/2-30, subtitle)  # Position the subtitle

    
    attacktrees = "Si-TaRA PDF"

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(A4)[0]/2, landscape(A4)[0]/2-180, projectname)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    

    c.setFont("Helvetica-Bold", title_font_size)
    c.drawCentredString(landscape(A4)[0]/2, landscape(A4)[0]/2-230, attacktrees)  # Position the title
    #c.drawCentredString(landscape(letter)[0]/2, landscape(letter)[0]/2-150, "Project Name: "+projectname)  # Position the title
    


    # c.setFont("Helvetica", 10)   
    # c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-210, taratype)
 
    #c.drawCentredString(landscape(letter)[0]/2+20, landscape(letter)[0]/2-180, taratype) 

    c.setFont("Helvetica", 10)
    c.drawCentredString(landscape(A4)[0]/2+300, landscape(A4)[0]/2-360, current_date)  #Position the date 

    c.setFont("Helvetica", manager_font_size)
    c.drawCentredString(landscape(A4)[0]/2-240, landscape(A4)[0]/2-360, "Security Manager Name: "+man_name)  #Position the Security Manager Name


    c.save()






def read_excel_file(file_path, sheets_to_read):
    xls = pd.ExcelFile(file_path)
    non_empty_data = {}
    for sheet_name, row_ranges in sheets_to_read.items():
        sheet_data = []
        for start_row, end_row in row_ranges:
            df = pd.read_excel(xls, sheet_name, skiprows=start_row - 1)
            if end_row:
                df = df.head(end_row - start_row + 1)  # Adjust dataframe rows if end_row is provided
            sheet_data.append(df.dropna(how='all').dropna(axis=1, how='all'))
        non_empty_data[sheet_name] = sheet_data
    return non_empty_data



 
# Custom function to create a paragraph with word wrap for each cell
def create_paragraphs(data_frame):
    style = ParagraphStyle(name='Normal', fontSize=10)
    return [[Paragraph(str(cell), style) if pd.notna(cell) else '' for cell in row] for row in data_frame.values]
 
# Function to write data to PDF in tabular format
def write_data_to_pdf(data, output_pdf_path):
    doc = SimpleDocTemplate(output_pdf_path, pagesize=landscape(letter), rightMargin=inch/2, leftMargin=inch/2)
    story = []
    styles = getSampleStyleSheet()
 
    i=0
    for sheet, df_list in data.items():
        for df in df_list:
    #for sheet, df in data.items():
            story.append(Paragraph(f'Sheet: {sheet}', styles['Heading2']))
            df.columns = ['' if col.startswith('Unnamed') else col for col in df.columns]

            #print(df.columns)
            #print("@@@@@")
            
            table_data = [df.columns.to_list()] + create_paragraphs(df)
        
            #print(table_data)
            #print("Close table_data")

            if i==0:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.lightblue),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])
            
            if i==1:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1),'#E5CB9F'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])
            
            if i==2:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#CCD4BF'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])

            if i==3:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#EEBAB2'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])

            if i==4:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#CFCFC4'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])

            if i==5:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#E7CBA9'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])

            if i==6:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#ababc9'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])


            if i==7:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#EBC8B4'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])


            if i==8:
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#A6AED0'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ])

            #print("Enter table_data")
            #print(table_data)
            #print("Closed table_data")
    
            table = Table(table_data, style=table_style, repeatRows=1, splitByRow=True)
            table.hAlign = 'LEFT'
            story.append(table)
            story.append(PageBreak())
        
        i=i+1
 
    doc.build(story)




def xl_to_pdf_LE(tara_file_name):
# Define your title and subtitle
    custom_title = "Si-TaRA"
    custom_subtitle = "Your Security Partner"
    #projectname = "Project Name"
    taratype = "Threat analysis and Risk Assessment"
    man_name = change["fullDocument"]["managerName"]

    # Output path for the custom title page
    custom_title_page_path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\custom_title_page.pdf'


    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\SiTaRA_Banner.PNG"

    #image_path = "C:\\Users\\UGY2KOR\\Desktop\\New_logo.PNG"
    
    # Create the custom title page
    create_title_page(custom_title, custom_subtitle, man_name,projectname,taratype,custom_title_page_path,image_path)

    sheets_to_read = {
    #'About':9,
    'TechDescription':[[4,6]],
    'Scope':[[4,5]],    
    'Assumptions': [[11, None]],  # Start from row 2 in Sheet1
    'MUCs' : [[32, None]],
    #'DSsConsequences':[[13,19]],
    'DSsConsequences': [[15,19],[22, None]],  # Start from row 5 in Sheet2
    'SecGoals': [[14, None]],  # Start from row 2 in Sheet1
    'ThreatsDSs' : [[16, None]],
    'ThreatEvaluation_LE': [[11,15],[17, None]],
    'RiskAssessment_LE': [[10, None]],  # Start from row 2 in Sheet1
    'SecurityNeeds_LE' : [[57, None]],
    'MngSummary_LE': [[26, None]],
    # Add more sheets as needed
    }


    data_to_pdf_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\table_output.pdf'

    output_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\final_output.pdf'

    # Replace 'your_excel_file.xlsx' with your actual Excel file path
    excel_data = read_excel_file(tara_file_name, sheets_to_read)
    write_data_to_pdf(excel_data,data_to_pdf_path)

    methodology_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\methodology_final.pdf'

    pdf_files=[custom_title_page_path,methodology_path,data_to_pdf_path]

    merge_multiple_pdfs(pdf_files, output_path)


    
    #input_file = 'C:\\Users\\UGY2KOR\\Desktop\\Final_PDF_Report_AP.pdf'  # Replace with your input PDF file
    watermark_file = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\watermark_text.pdf'  # Replace with your watermark PDF file
    output_file = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_LE.pdf'  # Replace with your output file name

    input_file = output_path
    
    # define the reader and writer objects
    reader_input = PdfReader(input_file)
    writer_output = PdfWriter()
    watermark_input = PdfReader(watermark_file)
    watermark = watermark_input.pages[0]
    
    # go through the pages one after the next
    for current_page in range(1,len(reader_input.pages)):
        merger = PageMerge(reader_input.pages[current_page])
        merger.add(watermark).render()
    
    # write the modified content to disk
    writer_output.write(output_file, reader_input)

    return output_file



def xl_to_pdf_AP(tara_file_name):

# Define your title and subtitle
    custom_title = "Si-TaRA"
    custom_subtitle = "Your Security Partner"
    #projectname = "Project Name"
    taratype = "Threat analysis and Risk Assessment"
    man_name = change["fullDocument"]["managerName"]

    # Output path for the custom title page
    custom_title_page_path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\custom_title_page.pdf'


    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    #image_path = "C:\\Users\\UGY2KOR\\Desktop\\New_logo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\SiTaRA_Banner.PNG"
    image_path1 = "C:\\Users\\Administrator\\Desktop\\Si-TaRA_front.png"

    # Create the custom title page
    create_title_page(custom_title, custom_subtitle, man_name,projectname,taratype,custom_title_page_path,image_path)

    # Define sheets and their starting rows here
    sheets_to_read = {
        #'About':9,
        'TechDescription':[[4,6]],
        'Scope':[[4,5]],
        'Assumptions': [[11, None]],  # Start from row 2 in Sheet1
        'MUCs' : [[32, None]],
        #'DSsConsequences':[[13,19]],
        'DSsConsequences': [[15,19],[22, None]],  # Start from row 5 in Sheet2
        'SecGoals': [[14, None]],  # Start from row 2 in Sheet1
        'ThreatsDSs' : [[16, None]],
        'ThreatEvaluation_AP': [[16,49],[51, None]],
        'RiskAssessment_AP': [[10, None]],  # Start from row 2 in Sheet1
        'SecurityNeeds_AP' : [[57, None]],
        'MngSummary_AP': [[26, None]],
        # Add more sheets as needed
    }


    data_to_pdf_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\table_output.pdf'

    output_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\final_output.pdf'

    # Replace 'your_excel_file.xlsx' with your actual Excel file path
    excel_data = read_excel_file(tara_file_name, sheets_to_read)
    write_data_to_pdf(excel_data,data_to_pdf_path)

    methodology_path='C:\\Users\\Administrator\\Desktop\\PDF_FILES\\methodology_final.pdf'

    pdf_files=[custom_title_page_path,methodology_path,data_to_pdf_path]

    merge_multiple_pdfs(pdf_files, output_path)


    
    #input_file = 'C:\\Users\\UGY2KOR\\Desktop\\Final_PDF_Report_AP.pdf'  # Replace with your input PDF file
    watermark_file = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\watermark_text.pdf'  # Replace with your watermark PDF file
    output_file = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf'  # Replace with your output file name

    input_file = output_path
    
    # define the reader and writer objects
    reader_input = PdfReader(input_file)
    writer_output = PdfWriter()
    watermark_input = PdfReader(watermark_file)
    watermark = watermark_input.pages[0]
    
    # go through the pages one after the next
    for current_page in range(1,len(reader_input.pages)):
        merger = PageMerge(reader_input.pages[current_page])
        merger.add(watermark).render()
    
    # write the modified content to disk
    writer_output.write(output_file, reader_input)

    return output_file



###########################################################################################################################



##################################################################################################################

def merge_multiple_pdfs1(pdf_paths, output_path):
    pdf_writer = PyPDF2.PdfWriter()

    for pdf_path in pdf_paths:
        pdf = PyPDF2.PdfReader(pdf_path)
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            pdf_writer.add_page(page)

    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

def create_title_page1(title, subtitle, man_name,projectname,taratype, output_path,image_path):
    c = canvas.Canvas(output_path, pagesize=A3)

    #background_color = HexColor("#0D1F2D")  # Dark blue background
    #background_color = HexColor("#ADD8E6")
    #text_color = HexColor("#FFFFFF")  # White text color

    #background_color = HexColor("#1A1A1A")  # Dark grayish-blue background
    #text_color = HexColor("#7FFF00")   # green radium
    #text_color = HexColor("#ADD8E6")   #light blue
    # Set the background color
    #current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    current_date = change["fullDocument"]["dateTime"]
    
    #background_color = HexColor("#330033")  # Dark grayish-blue background
    #text_color = HexColor("#00FFFF")

    background_color = HexColor("#FFFFFF")  # Dark grayish-blue background   #output 33
    text_color = HexColor("#330033")      #output 33

    #background_color = HexColor("#ADD8E6")  # Dark grayish-blue background   #output 34
    #text_color = HexColor("#330033")           #output 34


    

    c.setFillColor(background_color)
    c.rect(0, 0, A3[0], A3[1], fill=1)




    # Title
    c.setFillColor(text_color)

    #c.drawImage(image_path, 400, 600, width=30, height=50)

    c.drawImage(image_path, 600, 720, width=50, height=70)

    """
    # Title
    c.setFont("Helvetica-Bold", 45)
    c.drawCentredString(300, 600, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 10)
    c.drawCentredString(340, 640, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(320, 500, projectname)  # Position the title
    
    c.setFont("Helvetica", 10)
    c.drawCentredString(330, 480, taratype)    """


    c.setFont("Helvetica-Bold", 85)
    c.drawCentredString(A3[0]/2-20, A3[0]/2+300, title)  # Position the title

    # Subtitle
    c.setFont("Helvetica", 25)
    c.drawCentredString(A3[0]/2+40, A3[0]/2+370, subtitle)  # Position the subtitle

    c.setFont("Helvetica-Bold", 62)
    c.drawCentredString(A3[0]/2, A3[0]/2+100, projectname)  # Position the title
    
    c.setFont("Helvetica", 20)
    c.drawCentredString(A3[0]/2+20, A3[0]/2+70, taratype) 

    c.setFont("Helvetica", 20)
    c.drawCentredString(A3[0]/2+300, A3[0]/2-360, current_date)  #Position the date 

    c.setFont("Helvetica", 20)
    c.drawCentredString(A3[0]/2-240, A3[0]/2-360, man_name)  #Position the Security Manager Name

    c.save()



def xl_to_pdf_LE1(tara_file_name):
# Define your title and subtitle
    custom_title = "Si-TaRA"
    custom_subtitle = "Your Security Partner"
    #projectname = "Project Name"
    taratype = "Threat analysis and Risk Assessment"
    man_name = change["fullDocument"]["managerName"]

    # Output path for the custom title page
    custom_title_page_path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\custom_title_page.pdf'


    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    #image_path = "C:\\Users\\UGY2KOR\\Desktop\\New_logo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\banner.PNG"

    # Create the custom title page
    create_title_page(custom_title, custom_subtitle, man_name,projectname,taratype,custom_title_page_path,image_path)


    WB_PATH = tara_file_name
    PATH_TO_PDF1 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output11.pdf'
    PATH_TO_PDF2 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output12.pdf'
    PATH_TO_PDF3 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output13.pdf'
    PATH_TO_PDF4 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output14.pdf'
    PATH_TO_PDF5 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output15.pdf'
    PATH_TO_PDF6 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output16.pdf'
    PATH_TO_PDF7 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output17.pdf'
    PATH_TO_PDF8 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output18.pdf'
    PATH_TO_PDF9 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output19.pdf'
    PATH_TO_PDF_10 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output20.pdf'
    PATH_TO_PDF_11 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output21.pdf'
    PATH_TO_PDF_12 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output22.pdf'
    PATH_TO_PDF_13 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output23.pdf'
    PATH_TO_PDF_14 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output24.pdf'
    PATH_TO_PDF15 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Final_PDF_Report_LE.pdf'

    pdf_files = [custom_title_page_path,PATH_TO_PDF1,PATH_TO_PDF2,PATH_TO_PDF3,PATH_TO_PDF4,PATH_TO_PDF5,PATH_TO_PDF6,PATH_TO_PDF7,PATH_TO_PDF8,PATH_TO_PDF9,PATH_TO_PDF_10,PATH_TO_PDF_11,PATH_TO_PDF_12,PATH_TO_PDF_13]

    # Create a new instance of Excel without displaying it
    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False

    pdf_merger = PyPDF2.PdfMerger()

    try:
        
        i=0
        for pdf_file in pdf_files:
            if i==0:
                pdf_merger.append(pdf_file)
                i=i+1
                continue
            
            wb = excel.Workbooks.Open(WB_PATH)
            if(i==1):
                ws_source=wb.Worksheets("About")
            elif i==2:
                ws_source=wb.Worksheets("Methodology")
            elif i==3:
                ws_source=wb.Worksheets("TechDescription")
            elif i==4:
                ws_source=wb.Worksheets("Scope")
            elif i==5:
                ws_source=wb.Worksheets("Assumptions")
            elif i==6:
                ws_source=wb.Worksheets("MUCs")
            elif i==7:
                ws_source=wb.Worksheets("DSsConsequences")
            elif i==8:
                ws_source=wb.Worksheets("SecGoals")
            elif i==9:
                ws_source=wb.Worksheets("ThreatsDSs")
            elif i==10:
                ws_source=wb.Worksheets("ThreatEvaluation_LE")
            elif i==11:
                ws_source=wb.Worksheets("RiskAssessment_LE")
            elif i==12:
                ws_source=wb.Worksheets("SecurityNeeds_LE")
            elif i==13:
                ws_source=wb.Worksheets("MngSummary_LE")
            #else:
            #   ws_source = wb.Worksheets(i)  # Access the first worksheet, adjust the index as needed
            
            ws_source.PageSetup.Orientation = 1  # Set the page orientation
        # Adjust the page size to fit the content (for example, A3 size)
            ws_source.PageSetup.PaperSize = 8  # 8 corresponds to A3

        # Adjust the scaling to fit the content on a single page (100% scale)
            if i==2 or i==5 or i==6 or i==13:
                ws_source.PageSetup.Zoom = 50
            elif i==7 or i==8 or i==9:
                ws_source.PageSetup.Zoom = 30  
            elif i==12 or i==10:
                ws_source.PageSetup.Zoom = 35
            elif i==11:
                ws_source.PageSetup.Zoom = 40
            else:
                ws_source.PageSetup.Zoom = False


            
            ws_source.Select()
            
        # Export the selected worksheet to PDF
            wb.ActiveSheet.ExportAsFixedFormat(0, pdf_file)
            #print(i)
            
            #print(i)
        # Iterate through the PDF files and append them to the merger object
        #for pdf_file in pdf_files:
            pdf_merger.append(pdf_file)

        # Write the merged PDF to the output file
            i=i+1

            
        output_path = PATH_TO_PDF15

        merge_multiple_pdfs(pdf_files, output_path)

            

        #pdf_merger.write(PATH_TO_PDF15)
    except com_error as e:
        print(e)
    else:
        print('Succeeded.')
    finally:
        wb.Close(SaveChanges=False)  # Don't save changes to the Excel file
        excel.Quit()

    return PATH_TO_PDF15

    

def xl_to_pdf_AP1(tara_file_name):

# Define your title and subtitle
    custom_title = "Si-TaRA"
    custom_subtitle = "Your Security Partner"
    #projectname = "Project Name"
    taratype = "Threat analysis and Risk Assessment"
    man_name = change["fullDocument"]["managerName"]

    # Output path for the custom title page
    custom_title_page_path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\custom_title_page.pdf'


    image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    #image_path = "C:\\Users\\UGY2KOR\\Desktop\\New_logo.PNG"

    # Create the custom title page
    create_title_page(custom_title, custom_subtitle,man_name, projectname,taratype,custom_title_page_path,image_path)


    WB_PATH = tara_file_name
    PATH_TO_PDF1 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output11.pdf'
    PATH_TO_PDF2 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output12.pdf'
    PATH_TO_PDF3 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output13.pdf'
    PATH_TO_PDF4 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output14.pdf'
    PATH_TO_PDF5 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output15.pdf'
    PATH_TO_PDF6 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output16.pdf'
    PATH_TO_PDF7 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output17.pdf'
    PATH_TO_PDF8 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output18.pdf'
    PATH_TO_PDF9 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output19.pdf'
    PATH_TO_PDF10 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output20.pdf'
    PATH_TO_PDF11 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output21.pdf'
    PATH_TO_PDF12 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output22.pdf'
    PATH_TO_PDF13 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output23.pdf'
    PATH_TO_PDF14 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\output24.pdf'
    PATH_TO_PDF15 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Final_PDF_Report_AP.pdf'

    pdf_files = [custom_title_page_path,PATH_TO_PDF1,PATH_TO_PDF2,PATH_TO_PDF3,PATH_TO_PDF4,PATH_TO_PDF5,PATH_TO_PDF6,PATH_TO_PDF7,PATH_TO_PDF8,PATH_TO_PDF9,PATH_TO_PDF10,PATH_TO_PDF11,PATH_TO_PDF12,PATH_TO_PDF13,PATH_TO_PDF14]

    # Create a new instance of Excel without displaying it
    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False

    try:
        
        i=0
        for pdf_file in pdf_files:
            if i==0:
                pdf_merger = PyPDF2.PdfMerger()
                pdf_merger.append(pdf_file)
                i=i+1
                continue
            
            wb = excel.Workbooks.Open(WB_PATH)
            ws_source = wb.Worksheets(i)  # Access the first worksheet, adjust the index as needed

            ws_source.PageSetup.Orientation = 1  # Set the page orientation
        # Adjust the page size to fit the content (for example, A3 size)
            ws_source.PageSetup.PaperSize = 8  # 8 corresponds to A3

        # Adjust the scaling to fit the content on a single page (100% scale)
            if i==2 or i==5 or i==6 or i==12:
                ws_source.PageSetup.Zoom = 50
            elif i==7 or i==8 or i==9:
                ws_source.PageSetup.Zoom = 30
            elif i==10 or i==15:
                ws_source.PageSetup.Zoom = 30    
            elif i==13 or i==17:
                ws_source.PageSetup.Zoom = 40
            else:
                ws_source.PageSetup.Zoom = False
            #ws_source.PageSetup.FitToPagesTall = 1
            #ws_source.PageSetup.FitToPagesWide = 1
            ws_source.Select()

        # Export the selected worksheet to PDF
            wb.ActiveSheet.ExportAsFixedFormat(0, pdf_file)
            pdf_merger = PyPDF2.PdfMerger()

        # Iterate through the PDF files and append them to the merger object
        #for pdf_file in pdf_files:
            pdf_merger.append(pdf_file)

        # Write the merged PDF to the output file
            

            # print(i)
            i=i+1
            if i==11:
                i=i+1
                continue
            elif i==13:
                continue
            elif i==14:
                i=i+1
                continue
            elif i==16:
                i=i+1
                continue

        # List of paths to the PDF files to be merged and the output file
        #pdf_paths = ['file1.pdf', 'file2.pdf', 'file3.pdf']  # Add more file paths as needed
        output_path = PATH_TO_PDF15

        merge_multiple_pdfs(pdf_files, output_path)

            

        #pdf_merger.write(PATH_TO_PDF15)
    except com_error as e:
        print(e)
    else:
        print('Succeeded.')
    finally:
        wb.Close(SaveChanges=False)  # Don't save changes to the Excel file
        excel.Quit()
    
    return PATH_TO_PDF15

##########################################################################################################################

def generate_security_controls_doc(session_id, current_status):
    if (current_status == "processing"):
        iterate_interfaces_list()
        
        
        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})
        

        if (methodology == 'Attack Potential'):
            copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
        else:
            copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')

    
    
    
        security_query1 = f"SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[0].strip().lower()}`"
        print(security_query1)
        for i in range(0,len(interfaces_list_AP)):
            if i==0:
                continue
            else:
                security_query1 = security_query1+f" union all SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[i].strip().lower()}`"



        security_query2 = f"Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[0].strip().lower()}`"
        print(security_query2)
        for i in range(0,len(interfaces_list_AP)):
            if i==0:
                continue
            else:
                security_query2 = security_query2+f" union all Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[i].strip().lower()}`"




        global security_concepts_docx
        security_concepts_docx = create_security_concepts(tara_file_name, f"{TARA_FILES_PATH}{session_id}", methodology,security_query2)


        security_concepts_docx1 = header_footer3(security_concepts_docx)
        
        uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)


        
        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")

        for key in change["fullDocument"]["docType"]:
            if key == 'AP SCD' or key == 'LE SCD':
                shutil.copy(security_concepts_docx1,f"{ZIPFILES_PATH}{session_id}")
        
        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)


        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "generated"}})
        print("Process Completed")

################################################################################################################################

def update_tara_status_xl_secdoc(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)
    
    global counter_i

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        if methodology == 'Attack Potential':
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "processing"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
        else:
            print("2267", methodology)
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"lexl_status": "processing"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
                

        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})
        iterate_interfaces_list()
        if (methodology == 'Attack Potential'):
            copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
        else:
            copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')

        
        security_query1 = f"SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[0].strip().lower()}`"
                
        for i in range(0,len(interfaces_list_AP)):
                if i==0:
                    continue
                else:
                    security_query1 = security_query1+f" union all SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[i].strip().lower()}`"


        security_query2 = f"Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[0].strip().lower()}`"
        print(security_query2)
        for i in range(0,len(interfaces_list_AP)):
            if i==0:
                continue
            else:
                security_query2 = security_query2+f" union all Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[i].strip().lower()}`"

        global security_concepts_docx
        security_concepts_docx = create_security_concepts(tara_file_name, f"{TARA_FILES_PATH}{session_id}", methodology,security_query2)




        
        # print("The new security_concepts_file path")
        # print(security_concepts_docx)
        # global security_concepts_docx_path
        # security_concepts_docx_path = f"{TARA_FILES_PATH}{session_id}"
        # if (methodology == 'Attack Potential'):
        #     final_report = xl_to_pdf_AP(tara_file_name)
        # else:
        #     final_report = xl_to_pdf_LE(tara_file_name)



        security_concepts_docx1 = header_footer3(security_concepts_docx)
        
        uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)


        
        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")

        for key in change["fullDocument"]["docType"]:
            if key == 'AP .xlsm' or key == 'LE .xlsm': #.xlsm 
                if key == 'AP .xlsm':
                    print("2328","inside ap .xlsm")
                    shutil.copy(tara_file_name,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.xlsm")
                else:
                    shutil.copy(tara_file_name,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_LE.xlsm")


            if key == 'AP SCD' or key == 'LE SCD': # .xlsm scd
                shutil.copy(security_concepts_docx1,f"{ZIPFILES_PATH}{session_id}")



        



        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "Completed", "filePath": f'https://sitarastorage.s3.ap-south-1.amazonaws.com/Sitara_TARA_Storage/{email}/{session_id}/TRA_and_RRA_Template_v2.3.2.xlsm', "email": email}})
        # client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"filePath": tara_file_name}})
        # docCount = client.master_tara.UserData.find_one({"email": email})

        # print(docCount['docCount']) 

        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )

        # cursor = conn.cursor()
        # # updateUserData = DataBase("userdocdata")
        # global keys
        # for key in change["fullDocument"]["docType"]:
        #     updateUserDocQuery = f'insert into docsgenerated values("{projectname}", "{key}", "{change["fullDocument"]["dateTime"]}", "{email}")'
        #     print(updateUserDocQuery)
        #     cursor.execute(updateUserDocQuery)
        #     conn.commit()
        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)

        if methodology == 'Attack Potential':
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "generated"}})
        else:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"lexl_status": "generated"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "generated"}})

        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)
        

        print("Process Completed")
        




###########################################################################################################################################
        


def update_tara_status_VD(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)
    print(change["fullDocument"]["docType"])
    print("inside VDF")
    global counter_i

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        
        docty = change["fullDocument"]["docType"]

        
        
        if change["fullDocument"]["vulnDoc"]:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"VD_status": "processing"}})
        

        
        
        #client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})
        # iterate_interfaces_list()
        iterate_interfaces()
        # if ((methodology == 'Attack Potential') and ("AP .xlsm" in docty)):
        #     copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
        #     client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
        # else:
        #     copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')

        
        
        vulnerability_report(tmp_queue,f"{TARA_FILES_PATH}{session_id}")
        
        # print("The new security_concepts_file path")
        # print(security_concepts_docx)
        # global security_concepts_docx_path
        # security_concepts_docx_path = f"{TARA_FILES_PATH}{session_id}"
        # if (methodology == 'Attack Potential'):
        #     final_report = xl_to_pdf_AP(tara_file_name)
        # else:
        #     final_report = xl_to_pdf_LE(tara_file_name)


        # Folder path where PDFs are stored
        
        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")



        vulnerability_report_path=f"{TARA_FILES_PATH}{session_id}"+"\\vulnerability_report.xlsx"
        vulnerability_report_not_available=''

        if(vul_report_gen==True and os.path.exists(vulnerability_report_path)==False):
            vulnerability_report_not_available='\nNo vulnerability was found for the selected assests\n'





        # if(vul_report_gen==True and os.path.exists(vulnerability_report_path)):
        #     with open(vulnerability_report_path, 'rb') as attacktrees_attachment:
        #         vulnerability_attachment_content = attacktrees_attachment.read()
        # vulnerability_attachment=MIMEApplication(vulnerability_attachment_content, Name='vulnerability_document.xlsx')
        shutil.copy(vulnerability_report_path,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Vulnerability_Report.xlsx")
        # print('vulnerability report attached')
        # if methodology == "Attack Potential":
        #     AP_VD = client.master_tara.UserData.find_one({"email": email})
        #     client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"AP_VD": AP_VD["AP_VD"] + 1}})
        # else:
        #     LE_VD = client.master_tara.UserData.find_one({"email": email})
        #     client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"LE_VD": LE_VD["LE_VD"] + 1}})



        # sendEmail(tara_file_name, security_concepts_docx1,attacktreesfile1,tracefile1)
        # time.sleep(10)
        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "Completed", "filePath": f'https://sitarastorage.s3.ap-south-1.amazonaws.com/Sitara_TARA_Storage/{email}/{session_id}/TRA_and_RRA_Template_v2.3.2.xlsm', "email": email}})
        # client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"filePath": tara_file_name}})
        # docCount = client.master_tara.UserData.find_one({"email": email})

        # print(docCount['docCount']) 

        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )

        # cursor = conn.cursor()
        # # updateUserData = DataBase("userdocdata")
        # global keys
        # for key in change["fullDocument"]["docType"]:
        #     updateUserDocQuery = f'insert into docsgenerated values("{projectname}", "{key}", "{change["fullDocument"]["dateTime"]}", "{email}")'
        #     print(updateUserDocQuery)
        #     cursor.execute(updateUserDocQuery)
        #     conn.commit()
        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        
        
        
        
        
        
        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)
        
        if change["fullDocument"]["vulnDoc"]:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"VD_status": "generated"}})


################################################################################################################################################################################
        


        print("Process Completed")
        



#################################################################################################################################################################################










############################################################################################################################

def update_tara_status(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)
    
    global counter_i

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        
        docty = change["fullDocument"]["docType"]

        if 'AP .xlsm' in docty:
            print("inside if")
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "processing"}})
        if 'AP .PDF' in docty:
            print("inside AP .PDF")
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "processing"}})
        if change["fullDocument"]["vulnDoc"]:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"VD_status": "processing"}})
        if 'AP SCD' in docty or 'LE SCD' in docty:
            print("inside AP SCD" if 'AP SCD' in docty else "inside LE SCD")
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})

        if 'LE .xlsm' in docty:
            print('inside LE .xlsm')
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"lexl_status": "processing"}})
        
        #client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})
        iterate_interfaces_list()
        if (methodology == 'Attack Potential'):
            copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
        else:
            copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')

        
        security_query1 = f"SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[0].strip().lower()}`"
                
        for i in range(0,len(interfaces_list_AP)):
                if i==0:
                    continue
                else:
                    security_query1 = security_query1+f" union all SELECT Security_Controls FROM Security_Controls.`{interfaces_list_AP[i].strip().lower()}`"


        security_query2 = f"Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[0].strip().lower()}`"
        print(security_query2)
        for i in range(0,len(interfaces_list_AP)):
            if i==0:
                continue
            else:
                security_query2 = security_query2+f" union all Select Security_Control, Security_Control_Description from Security_Goals_Controls.`{interfaces_list_AP[i].strip().lower()}`"            


        global security_concepts_docx
        security_concepts_docx = create_security_concepts(tara_file_name, f"{TARA_FILES_PATH}{session_id}", methodology,security_query2)
        vulnerability_report(tmp_queue,f"{TARA_FILES_PATH}{session_id}")
        
        # print("The new security_concepts_file path")
        # print(security_concepts_docx)
        # global security_concepts_docx_path
        # security_concepts_docx_path = f"{TARA_FILES_PATH}{session_id}"
        # if (methodology == 'Attack Potential'):
        #     final_report = xl_to_pdf_AP(tara_file_name)
        # else:
        #     final_report = xl_to_pdf_LE(tara_file_name)


        # Folder path where PDFs are stored
        folder_path = 'C:\\Users\\Administrator\\Desktop\\Attack_Trees_folder'


        trace_fd_path = 'C:\\Users\\Administrator\\Desktop\\Tracebility_Graphs_Folder'



        # Output path for the combined PDF
        global output_path

        global traceout_path

        ATTACKTREES_PATH= 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\'


        Tracebility_PATH = 'C:\\Users\\Administrator\\Desktop\\TRACE_FILES\\'


        directory_path2 = f"{Tracebility_PATH}{session_id}"

        if not os.path.exists(directory_path2):
            os.mkdir(directory_path2)
            # print(f"Directory '{directory_path2}' created successfully.")
        else:
            print(f"Directory '{directory_path2}' already exists.")   





        directory_path1 = f"{ATTACKTREES_PATH}{session_id}"

        if not os.path.exists(directory_path1):
            os.mkdir(directory_path1)
            print(f"Directory '{directory_path1}' created successfully.")
        else:
            print(f"Directory '{directory_path1}' already exists.")   


        #os.mkdir(f"{ATTACKTREES_PATH}{session_id}")

        output_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Attack_Trees.pdf'
                            

        traceout_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Tracebility_graphs.pdf'


        global attacktreesfile

        global tracefile
        
        attacktreesfile = merge_pdfs_from_folder(folder_path, interfaces_list_AP, output_path)
        # print(attacktreesfile)


        tracefile = merge_pdfs_from_folder1(trace_fd_path, interfaces_list_AP, traceout_path)
        # print(tracefile)

        
        # print(interfaces_list_AP)


        #final_report1 = header_footer_AP_LE(final_report)

        attacktreesfile1 = header_footer1(attacktreesfile)

        


        tracefile1 = header_footer2(tracefile)

        security_concepts_docx1 = header_footer3(security_concepts_docx)
        
        uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)

        sendEmail(tara_file_name, security_concepts_docx1,attacktreesfile1,tracefile1)
        client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})

        time.sleep(10)
        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "Completed", "filePath": f'https://sitarastorage.s3.ap-south-1.amazonaws.com/Sitara_TARA_Storage/{email}/{session_id}/TRA_and_RRA_Template_v2.3.2.xlsm', "email": email}})
        # client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"filePath": tara_file_name}})
        # docCount = client.master_tara.UserData.find_one({"email": email})

        # print(docCount['docCount']) 

        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )

        # cursor = conn.cursor()
        # # updateUserData = DataBase("userdocdata")
        # global keys
        # for key in change["fullDocument"]["docType"]:
        #     updateUserDocQuery = f'insert into docsgenerated values("{projectname}", "{key}", "{change["fullDocument"]["dateTime"]}", "{email}")'
        #     print(updateUserDocQuery)
        #     cursor.execute(updateUserDocQuery)
        #     conn.commit()
        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)
        

        if 'AP .xlsm' in docty:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
        if 'AP .PDF' in docty:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "generated"}})
        if change["fullDocument"]["vulnDoc"]:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"VD_status": "generated"}})
        if 'AP SCD' in docty:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "generated"}})


################################################################################################################################################################################
        if 'LE .xlsm' in docty:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"lexl_status": "generated"}})


        print("Process Completed")
        



#################################################################################################################################################################################






def update_tara_status_xl(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)
    
    global counter_i

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        
        client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "processing"}})

        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})
        iterate_interfaces_list()
        if (methodology == 'Attack Potential'):
            copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
        else:
            copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')


        if (methodology == 'Attack Potential'):
            final_report = xl_to_pdf_AP(tara_file_name)
        else:
            final_report = xl_to_pdf_LE(tara_file_name)




        final_report1 = header_footer_AP_LE(final_report)


        



        uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)

        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "Completed", "filePath": f'https://sitarastorage.s3.ap-south-1.amazonaws.com/Sitara_TARA_Storage/{email}/{session_id}/TRA_and_RRA_Template_v2.3.2.xlsm', "email": email}})
        # client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"filePath": tara_file_name}})
        # docCount = client.master_tara.UserData.find_one({"email": email})

        # print(docCount['docCount']) 

        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )


        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")   
        shutil.copy(tara_file_name,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.xlsm")
        

        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)



        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)           
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
        print("Process Completed")




######################################################################################################################################################
        



def update_tara_status_xl_pdf(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)
    
    global counter_i

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        
        client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "processing"}})

        client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "processing"}})
        

        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "processing", "filePath": 'generating'}})

        


        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")   
        


        # Folder path where PDFs are stored
        folder_path = 'C:\\Users\\Administrator\\Desktop\\Attack_Trees_folder'


        trace_fd_path = 'C:\\Users\\Administrator\\Desktop\\Tracebility_Graphs_Folder'



        # Output path for the combined PDF
        global output_path

        global traceout_path

        ATTACKTREES_PATH= 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\'


        Tracebility_PATH = 'C:\\Users\\Administrator\\Desktop\\TRACE_FILES\\'


        directory_path2 = f"{Tracebility_PATH}{session_id}"

        if not os.path.exists(directory_path2):
            os.mkdir(directory_path2)
            # print(f"Directory '{directory_path2}' created successfully.")
        else:
            print(f"Directory '{directory_path2}' already exists.")   





        directory_path1 = f"{ATTACKTREES_PATH}{session_id}"

        if not os.path.exists(directory_path1):
            os.mkdir(directory_path1)
            print(f"Directory '{directory_path1}' created successfully.")
        else:
            print(f"Directory '{directory_path1}' already exists.")   


        #os.mkdir(f"{ATTACKTREES_PATH}{session_id}")

        output_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Attack_Trees.pdf'
                            

        traceout_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Tracebility_graphs.pdf'


        global attacktreesfile

        global tracefile
        iterate_interfaces()
        print(interfaces_list_AP1)
        attacktreesfile = merge_pdfs_from_folder(folder_path, interfaces_list_AP1, output_path)
        # print(attacktreesfile)


        tracefile = merge_pdfs_from_folder1(trace_fd_path, interfaces_list_AP1, traceout_path)
        # print(tracefile)

        
        # print(interfaces_list_AP)


        #final_report1 = header_footer_AP_LE(final_report)

        attacktreesfile1 = header_footer1(attacktreesfile)

        


        tracefile1 = header_footer2(tracefile)

        
        # uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)
        # uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        # uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        
        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)#
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )

        # cursor = conn.cursor()
        # # updateUserData = DataBase("userdocdata")
        # global keys
        # for key in change["fullDocument"]["docType"]:
        #     updateUserDocQuery = f'insert into docsgenerated values("{projectname}", "{key}", "{change["fullDocument"]["dateTime"]}", "{email}")'
        #     print(updateUserDocQuery)
        #     cursor.execute(updateUserDocQuery)
        #     conn.commit()
        


        outputpdf = generatePDF(attacktreesfile1,tracefile1)


        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        print("PDF generated")
        client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "generated"}})








        iterate_interfaces_list()
        if (methodology == 'Attack Potential'):
            copy_sheets(AP_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_AP')
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
        else:
            copy_sheets(LE_Sheet_Copy, tara_file_name,projectname,scope,techdescription,'MngSummary_LE')


        shutil.copy(tara_file_name,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.xlsm")
        

        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)

        



   #     uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)

        client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"status": "Completed", "filePath": f'https://sitarastorage.s3.ap-south-1.amazonaws.com/Sitara_TARA_Storage/{email}/{session_id}/TRA_and_RRA_Template_v2.3.2.xlsm', "email": email}})
        # client.master_tara.Interfaces.find_one_and_update({"sessionId": session_id}, {"$set": {"filePath": tara_file_name}})
        # docCount = client.master_tara.UserData.find_one({"email": email})

        # print(docCount['docCount']) 

        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )


        





        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)           
            # client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"xl_status": "generated"}})
            
    
        print("Process Completed")

#######################################################SECURITY CONCEPTS###############################################################################





def update_tara_status_new(session_id, current_status):
    """
    Update the status field in the record with the given session ID.
    If the status is "processing", call the iterate_interfaces_list function.
    """
    # print("Inside Update Data Base- and Current Status outside IF condition is",session_id, current_status)

    global counter_i
    client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "processing"}})

    if (current_status == "processing"):
        # print("current_status in processing is", current_status)
        iterate_interfaces()    
        
        # Folder path where PDFs are stored
        folder_path = 'C:\\Users\\Administrator\\Desktop\\Attack_Trees_folder'


        trace_fd_path = 'C:\\Users\\Administrator\\Desktop\\Tracebility_Graphs_Folder'



        # Output path for the combined PDF
        global output_path

        global traceout_path

        ATTACKTREES_PATH= 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\'


        Tracebility_PATH = 'C:\\Users\\Administrator\\Desktop\\TRACE_FILES\\'


        directory_path2 = f"{Tracebility_PATH}{session_id}"

        if not os.path.exists(directory_path2):
            os.mkdir(directory_path2)
            # print(f"Directory '{directory_path2}' created successfully.")
        else:
            print(f"Directory '{directory_path2}' already exists.")   





        directory_path1 = f"{ATTACKTREES_PATH}{session_id}"

        if not os.path.exists(directory_path1):
            os.mkdir(directory_path1)
            print(f"Directory '{directory_path1}' created successfully.")
        else:
            print(f"Directory '{directory_path1}' already exists.")   


        #os.mkdir(f"{ATTACKTREES_PATH}{session_id}")

        output_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Attack_Trees.pdf'
                            

        traceout_path = f'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\{session_id}\\Tracebility_graphs.pdf'


        global attacktreesfile

        global tracefile
        print(interfaces_list_AP1)
        attacktreesfile = merge_pdfs_from_folder(folder_path, interfaces_list_AP1, output_path)
        # print(attacktreesfile)


        tracefile = merge_pdfs_from_folder1(trace_fd_path, interfaces_list_AP1, traceout_path)
        # print(tracefile)

        
        # print(interfaces_list_AP)


        #final_report1 = header_footer_AP_LE(final_report)

        attacktreesfile1 = header_footer1(attacktreesfile)

        


        tracefile1 = header_footer2(tracefile)

        
        # uploadDirectory('C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\','sitarastorage', session_id)
        # uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        # uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)
        
        counter_i = 0
        #replace_sheets(LE_Sheet_Copy, tara_file_name, 3)#
        # if (methodology == 'Attack Potential'):
        #     print("src path is ", AP_Sheet_Copy)
        # else:
        #     print("src path is ", LE_Sheet_Copy)
            
        # print("dest file is", tara_file_name)
        #getFiles(LE_Sheet_Copy, tara_file_name)
        # conn = mysql.connector.connect(
        # host="tara-database.cpwdcswc6fon.ap-south-1.rds.amazonaws.com",
        # user="admin",
        # password="phantom2022*",
        #  #database="tara_new"
        # database="userdocdata"
        # )

        # cursor = conn.cursor()
        # # updateUserData = DataBase("userdocdata")
        # global keys
        # for key in change["fullDocument"]["docType"]:
        #     updateUserDocQuery = f'insert into docsgenerated values("{projectname}", "{key}", "{change["fullDocument"]["dateTime"]}", "{email}")'
        #     print(updateUserDocQuery)
        #     cursor.execute(updateUserDocQuery)
        #     conn.commit()
        


        outputpdf = generatePDF(attacktreesfile1,tracefile1)
        print("Entering critical area 0")

        uploadFileToS3(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)

        print("Entering critical area intermission")

        uploadDirectory(f'C:\\Users\\Administrator\\Documents\\ZIP_FILES\\{session_id}', 'sitarastorage', session_id)

        print("Entering critical area 1")
        for key in change["fullDocument"]["docType"]:
            UserDocStatusQuery = {
                "projectName": projectname,
                "docType": key,
                "timeStamp": change["fullDocument"]["dateTime"],
                "email": email
            }
            print("Entering critical area 2")
            client.master_tara.UserDocData.insert_one(UserDocStatusQuery)
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "generated"}})
        print("Process Completed")







def create_folder_template(session_id):
    """
    Create a new folder with the given session ID and copy the TRA and RRA template to the folder.
    """
    global tara_file_name
    # print("Inside Create Folder")
    # create a new folder with the given session ID

    directory_path = f"{TARA_FILES_PATH}{session_id}"

    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
        # print(f"Directory '{directory_path}' created successfully.")
    else:
        print(f"Directory '{directory_path}' already exists.")    

    #os.mkdir(f"{TARA_FILES_PATH}{session_id}")
    # copy the TRA and RRA template to the folder
    #file_path = f"{TARA_FILES_PATH}{session_id}"
    global methodology 
    methodology = change["fullDocument"]["methodology"]
    if (methodology == 'Attack Potential'):
        # print("if is true")
        tara_file_name = shutil.copy(TRA_RRA_TEMPLATE_AP, f"{TARA_FILES_PATH}{session_id}")
    else:
        # print("in else statement")
        tara_file_name = shutil.copy(TRA_RRA_TEMPLATE_LE, f"{TARA_FILES_PATH}{session_id}")
    # delete_sheets(tara_file_name)
    # print("File Path is ",tara_file_name)
    #delete_sheets(tara_file_name)
    return tara_file_name

def generatePDF(attack_trees_path,trace_path):
    pdf_data = client.tara_pdf.PDF.find_one({'filename': session_id})
    if pdf_data:
        pdf_content = pdf_data['pdf']


                    # print(pdf_content,"This is pdf_content")

                #file = Response(pdf_content, content_type='application/pdf')
        path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf' 


        try:
            with open(path,'wb') as f:
                f.write(pdf_content)
        except Exception as e:
            print(e)


        newpath = header_footer4(path)

                
                # # Serve the PDF as a response with Content-Type set to application/pdf
                
        path1 = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP1.pdf'

        finalpath = merge_pdfs_from_folder2(attack_trees_path,trace_path,path1)
        print(finalpath, "finalpath")

        ZIPFILES_PATH = "C:\\Users\\Administrator\\Documents\\ZIP_FILES\\"



        zip_dir_path = f"{ZIPFILES_PATH}{session_id}"

        if not os.path.exists(zip_dir_path):
            os.mkdir(zip_dir_path)
            print(f"Directory '{zip_dir_path}' created successfully.")
        else:
            print(f"Directory '{zip_dir_path}' already exists.")   
        shutil.copy(finalpath,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf")
        # client.tara_pdf.PDF.find_one_and_update({})
                   # shutil.copy(output_zip_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Attack_Trees.7z")
                   # shutil.copy(output_zip_trace_file,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Traceability_Graphs.7z")
                

                #shutil.copy(final_pdf_report,f"{ZIPFILES_PATH}{session_id}\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf")
        return finalpath        
                
        
    
def get_new_record(session_id):
    """
    Call the create_folder_template function to create a new folder and copy the TRA and RRA template.
    Update the status of the record to "processing" and sleep for 10 seconds.
    Update the status of the record to "Completed".
    """
    # call the create_folder_template function to create a new folder and copy the TRA and RRA template
    # if change["fullDocument"]["docType"] == "AP .xlsm" or change["fullDocument"]["docType"] == "AP .xlsm ":
    create_folder_template(session_id)
    print(change["fullDocument"]["vulnDoc"], "docType")

    doccount = change["fullDocument"]["docType"]
    print(doccount)

    vdstat = change["fullDocument"]["vulnDoc"]

    pdfxlsm = ['AP .xlsm', 'AP .PDF']
    pdfxlsm1 = ['AP .PDF', 'AP .xlsm']

    pdf = ["AP .PDF"]

    scd = ["AP SCD"]

    LESCD = ["LE SCD"]

    pdfscd = ['AP .PDF', 'AP SCD']
    scdpdf = ['AP SCD', 'AP .PDF']


    xlsmsecdoc = ["AP .xlsm", "AP SCD"]
    xlsmsecdoc1 = ["AP SCD", "AP .xlsm"]

    lexlsmsecdoc = ["LE .xlsm", "LE SCD"]
    lexlsmsecdoc1 = ["LE SCD", "LE .xlsm"]

    xlsmvul = ["AP .xlsm", "AP VD"]
    xlsmvul1 = ["AP VD", "AP .xlsm"]





    pdfScdXlsmVd = ["AP .PDF", "AP .xlsm", "AP SCD"]
    xlsmpdfscd = ['AP .xlsm', 'AP .PDF', 'AP SCD']
    scdxlsmpdf = ['AP .SCD', 'AP .xlsm', 'AP PDF']

   # for key in change["fullDocument"]["docType"]:
     
    if((doccount == pdfxlsm) or (doccount == pdfxlsm1) or (doccount == pdfxlsm1 and vdstat) or (doccount == pdfxlsm and vdstat)):
        print("inside pdfxl")
        if vdstat:
            update_tara_status_VD(session_id,"processing")
            update_tara_status_xl_pdf(session_id,"processing")
        else:
            update_tara_status_xl_pdf(session_id,"processing")
    elif((doccount == pdfScdXlsmVd and vdstat) or (doccount == xlsmpdfscd and vdstat) or (doccount == scdxlsmpdf and vdstat)):
        print("inside pdfscdxlscdvd")
        update_tara_status(session_id,"processing")
    elif((doccount == pdf)):
        print("inside pdf")
        if vdstat:
            update_tara_status_new(session_id,"processing")
            update_tara_status_VD(session_id,"processing")
        else:
            update_tara_status_new(session_id,"processing")
    elif((doccount == xlsmsecdoc) or (doccount == xlsmsecdoc1) or (doccount == lexlsmsecdoc) or (doccount == lexlsmsecdoc1) or (doccount == xlsmsecdoc and vdstat) or (doccount == xlsmsecdoc1 and vdstat)):
        if vdstat:
            update_tara_status_VD(session_id,"processing")
            update_tara_status_xl_secdoc(session_id,"processing")
        else:
            update_tara_status_xl_secdoc(session_id,"processing")
    elif((doccount == scd) or (doccount == LESCD) or (doccount == scd and vdstat) or (doccount == LESCD and vdstat)):
        print("inside onlysec")
        if vdstat:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
            update_tara_status_VD(session_id,"processing")
            generate_security_controls_doc(session_id, "processing")
        else:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
            generate_security_controls_doc(session_id, "processing")
    elif((doccount == scdpdf) or (doccount == pdfscd) or (doccount == scdpdf and vdstat) or (doccount == pdfscd and vdstat)):
        if vdstat:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "processing"}})   
            update_tara_status_new(session_id,"processing")
            update_tara_status_VD(session_id,"processing")
            generate_security_controls_doc(session_id, "processing")
        else:
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"secdoc_status": "processing"}})
            client.master_tara.UserData.find_one_and_update({"email": email}, {"$set": {"PDF_status": "processing"}})   
            update_tara_status_new(session_id,"processing")
            generate_security_controls_doc(session_id, "processing")
    elif(vdstat):
        update_tara_status_VD(session_id,"processing")
        
    else:
        print("inside default")
        update_tara_status(session_id, "processing")
    


    # #delete_workbook_sheets(new_file)
  
    # # print("Value in the Job Queue is ", job_queue)
    # # update the status of the record to "processing"
    #     elif key == "AP .xlsm":
    #         update_tara_status_xl(session_id, "processing")

    #     elif key == "AP .xlsm AP .PDF":
    #         update_tara_status_new(session_id,"processing")
    #         update_tara_status_xl(session_id, "processing")
    # # sleep for 10 seconds

    #     else:
            
    
    #update_tara_status(session_id,"processing")
    #update_tara_status(session_id,"processing")
    # update the status of the record to "Completed"
   #
    



def iterate_interfaces():
    

    if(len(job_queue) >=0):
        job_queue.clear()
        
    global projectname
    global cloud_assets
    global interfaces_assets
    global database_assets
    global application_assets
    global vul_report_gen
    global scope
    global techdescription
    global automotives_assets
    global astArch


    aws_AssetLen = change["fullDocument"]["cloud"]
    cloud_assets=aws_AssetLen
    # azure_AssetLen = change["fullDocument"]["azureAssets"]
    device_AssetLen = change["fullDocument"]["device"]
    interfaces_AssetLen = change["fullDocument"]["interfaces"]
    interfaces_assets=interfaces_AssetLen
    db_AssetLen = change["fullDocument"]["DBAssets"]
    database_assets=db_AssetLen
    application_AssetLen = change["fullDocument"]["application"]
    application_assets=application_AssetLen
    projectname=change["fullDocument"]["projectName"]
    vul_report_gen=change["fullDocument"]["vulnDoc"]
    scope=change["fullDocument"]["scope"]
    techdescription= change["fullDocument"]["techDesc"]
    automotives_assets= change["fullDocument"]["automotive"]

    astArch = change["fullDocument"]["astArch"]

    # mobileapplication_AssetLen = change["fullDocument"]["mobileapplication"]
    
    if(len(aws_AssetLen) > 0):
        job_queue.append(aws_AssetLen)
    
    if(len(interfaces_AssetLen) > 0):
        job_queue.append(interfaces_AssetLen)
        
    #if(len(azure_AssetLen) > 0):
        #job_queue.append(azure_AssetLen)
        
    if(len(db_AssetLen) > 0):
        job_queue.append(db_AssetLen)

    if(len(application_AssetLen) > 0):
       job_queue.append(application_AssetLen)

    if(len(automotives_assets) > 0):
        job_queue.append(automotives_assets)


    if(len(astArch) > 0):
        job_queue.append(astArch)


    # if(len(mobileapplication_AssetLen) > 0):
    #    job_queue.append(mobileapplication_AssetLen)


    # print("Methodology used is", methodology)
                       
    #print("awsAssets  is ", len(change["fullDocument"]["awsAssets"]))
    # append the interfaces field from the change document to the job_queue list
    # job_queue.append(change["fullDocument"]["interfaces"])
    
    # insert the session ID at the beginning of the job_queue list
    job_queue.insert(0, change["fullDocument"]["sessionId"])
    
    # print("length of job_queue is ", job_queue)
    #tmp_queue = []
    # split the string on the commas to get a list of interface names
    #tmp_queue = str(job_queue).split(",")
    # filter out any empty strings in the list
    #tmp_queue = [x for x in tmp_queue if x]
    global tmp_queue
    tmp_queue = []
    # if(len(tmp_queue) >= 0):
    #     tmp_queue.clear()
    
    string = ""
    for i in range(len(str(job_queue))):
        if str(job_queue)[i].isalpha() or str(job_queue)[i] == '@' or str(job_queue)[i].isnumeric() or str(job_queue)[i] == ' ' or str(job_queue)[i] == '-' or str(job_queue)[i] == '&' or str(job_queue)[i] == '_' or str(job_queue)[i] == '(' or str(job_queue)[i] == ')':
            string=string+str(job_queue)[i]
        elif str(job_queue)[i]=="," or str(job_queue)[i]=="]":
            tmp_queue.append(string)
            string=""
            

    while("" in tmp_queue):
        tmp_queue.remove("")
    

    tmp_queue.pop(0)
    
    

    tmp_queue = [interface.strip().lower() for interface in tmp_queue]

    
    global interfaces_list_AP1
    interfaces_list_AP1 = tmp_queue

    print(interfaces_list_AP1)

    return interfaces_list_AP1

def iterate_interfaces_list():
    """
    Append the interfaces field from the change document to the job_queue list and insert the session ID at the beginning of the list.
    Split the string on the commas to get a list of interface names.
    Filter out any empty strings in the list and pass it to the build_tara_query function.
    """
    # global change

    # print("inside iterate interfaces")
    
    if(len(job_queue) >=0):
        job_queue.clear()
        
    global projectname
    global cloud_assets
    global interfaces_assets
    global database_assets
    global application_assets
    global vul_report_gen
    global scope
    global techdescription
    global automotives_assets
    global astArch



    aws_AssetLen = change["fullDocument"]["cloud"]
    cloud_assets=aws_AssetLen
    # azure_AssetLen = change["fullDocument"]["azureAssets"]
    device_AssetLen = change["fullDocument"]["device"]
    interfaces_AssetLen = change["fullDocument"]["interfaces"]
    interfaces_assets=interfaces_AssetLen
    db_AssetLen = change["fullDocument"]["DBAssets"]
    database_assets=db_AssetLen
    application_AssetLen = change["fullDocument"]["application"]
    application_assets=application_AssetLen
    projectname=change["fullDocument"]["projectName"]
    vul_report_gen=change["fullDocument"]["vulnDoc"]
    scope=change["fullDocument"]["scope"]
    techdescription= change["fullDocument"]["techDesc"]
    automotives_assets= change["fullDocument"]["automotive"]

    astArch = change["fullDocument"]["astArch"]


    # mobileapplication_AssetLen = change["fullDocument"]["mobileapplication"]
    
    if(len(aws_AssetLen) > 0):
        job_queue.append(aws_AssetLen)
    
    if(len(interfaces_AssetLen) > 0):
        job_queue.append(interfaces_AssetLen)
        
    #if(len(azure_AssetLen) > 0):
        #job_queue.append(azure_AssetLen)
        
    if(len(db_AssetLen) > 0):
        job_queue.append(db_AssetLen)

    if(len(application_AssetLen) > 0):
       job_queue.append(application_AssetLen)

    if(len(automotives_assets) > 0):
        job_queue.append(automotives_assets)

    if(len(astArch) > 0):
        job_queue.append(astArch)

    # if(len(mobileapplication_AssetLen) > 0):
    #    job_queue.append(mobileapplication_AssetLen)


    # print("Methodology used is", methodology)
                       
    #print("awsAssets  is ", len(change["fullDocument"]["awsAssets"]))
    # append the interfaces field from the change document to the job_queue list
    # job_queue.append(change["fullDocument"]["interfaces"])
    
    # insert the session ID at the beginning of the job_queue list
    job_queue.insert(0, change["fullDocument"]["sessionId"])
    
    # print("length of job_queue is ", job_queue)
    #tmp_queue = []
    # split the string on the commas to get a list of interface names
    #tmp_queue = str(job_queue).split(",")
    # filter out any empty strings in the list
    #tmp_queue = [x for x in tmp_queue if x]
    global tmp_queue
    tmp_queue = []
    # if(len(tmp_queue) >= 0):
    #     tmp_queue.clear()
    
    string = ""
    for i in range(len(str(job_queue))):
        if str(job_queue)[i].isalpha() or str(job_queue)[i] == '@' or str(job_queue)[i].isnumeric() or str(job_queue)[i] == ' ' or str(job_queue)[i] == '-' or str(job_queue)[i] == '&' or str(job_queue)[i] == '_' or str(job_queue)[i] == '(' or str(job_queue)[i] == ')':
            string=string+str(job_queue)[i]
        elif str(job_queue)[i]=="," or str(job_queue)[i]=="]":
            tmp_queue.append(string)
            string=""

            

    while("" in tmp_queue):
        tmp_queue.remove("")
    # print("Venkat list",tmp_queue)

    # pass the list of interface names to the build_tara_query function
    build_tara_query(tmp_queue)

def build_tara_query(interfaces_list):
    """
    Build a query string using the interface names in the interfaces_list.
    Pass the query string to the get_tara_results function.
    """
    # create an empty query string
    query = ""
    # iterate through the interface names in the interfaces_list
    var1=""
    var2 = ""
    global assets_list
    interfaces_list.pop(0)
    assets_list = interfaces_list

    

    interfaces_list = [interface.strip().lower() for interface in interfaces_list]

    
    global interfaces_list_AP
    interfaces_list_AP = interfaces_list

    #print(interfaces_list)

    for i in range(0,len(interfaces_list)):
        #if var1=="":
        #    var1=tmp_queue[i].lower()+" c"+str(i)
            # new_str=tmp_queue[i]
            # var2="c"+str(i)+".*"
        #else:
        #    var1=var1+","+tmp_queue[i].lower()+" c"+str(i)
        #    # new_str=tmp_queue[i]
            # x=new_str[0].lower()
        #    var2=var2+","+"c"+str(i)+".*"

        #print(interfaces_list.strip().lower())

        if i==0:
            var1 = f"select * from `{interfaces_list[i].strip().lower()}`"
        else:
            var1 = var1+f" union all select * from `{interfaces_list[i].strip().lower()}`"



    # for i in range(0,len(interfaces_list)):
        
    #     if i==0:
    #         var2 = f"select `Asset Description`,`Security Need Description_Threats` from `{interfaces_list[i].strip().lower()}`"
    #     else:
    #         var2 = var2+f" union all select `Asset Description`,`Security Need Description_Threats` from `{interfaces_list[i].strip().lower()}`"       

    # print(var2)

    # pass the query string to the get_tara_results function
    get_tara_results(var1)


#####################################################################################################################


def merge_pdfs_from_folder(folder_path, pdf_filenames, output_path):
    pdf_writer = PyPDF2.PdfWriter()

    custom_title_path_AP ='C:\\Users\\Administrator\\Desktop\\Attack_Trees_folder\\custom_title_page_AP.pdf'

    man_name = change["fullDocument"]["managerName"]



    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\SiTaRA_Banner.PNG"

    create_title_page_AP(man_name, projectname,custom_title_path_AP,image_path)


    extra_pdf = PyPDF2.PdfReader(custom_title_path_AP)
    for page_num in range(len(extra_pdf.pages)):
        page = extra_pdf.pages[page_num]
        pdf_writer.add_page(page)
    
    #pdf_path=os.path.join(extra_pdf)
    for filename in pdf_filenames:
        pdf_name_with_extension = filename + '.pdf'  # Append '.pdf' to the filename
        pdf_path = os.path.join(folder_path, pdf_name_with_extension)
        if os.path.exists(pdf_path):
            pdf = PyPDF2.PdfReader(pdf_path)
            for page_num in range(len(pdf.pages)):
                page = pdf.pages[page_num]
                pdf_writer.add_page(page)
        else:
            # print(f"PDF '{pdf_name_with_extension}' does not exist in the folder.")
            continue

    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

    # print("output_path is", output_path)
    return output_path


def merge_pdfs_from_folder1(folder_path, pdf_filenames, output_path):
    pdf_writer = PyPDF2.PdfWriter()


    custom_title_path_TR ='C:\\Users\\Administrator\\Desktop\\Tracebility_Graphs_Folder\\custom_title_page_TR.pdf'

    man_name = change["fullDocument"]["managerName"]



    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\SiTaRA_Banner.PNG"

    create_title_page_TR(man_name, projectname,custom_title_path_TR,image_path)


    extra_pdf = PyPDF2.PdfReader(custom_title_path_TR)
    for page_num in range(len(extra_pdf.pages)):
        page = extra_pdf.pages[page_num]
        pdf_writer.add_page(page)
    

    
    #pdf_path=os.path.join(extra_pdf)
    for filename in pdf_filenames:
        pdf_name_with_extension = filename + '.pdf'  # Append '.pdf' to the filename
        pdf_path = os.path.join(folder_path, pdf_name_with_extension)
        if os.path.exists(pdf_path):
            pdf = PyPDF2.PdfReader(pdf_path)
            for page_num in range(len(pdf.pages)):
                page = pdf.pages[page_num]
                pdf_writer.add_page(page)
        else:
            # print(f"PDF '{pdf_name_with_extension}' does not exist in the folder.")
            continue

    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

    # print("output_path is", output_path)
    return output_path





def merge_pdfs_from_folder2(attacktrees,traceabilitypath,output_path):
    pdf_writer = PyPDF2.PdfWriter()


    custom_title_path_TaRA ='C:\\Users\\Administrator\\Desktop\\TARA_PDF\\custom_title_page_TaRA.pdf'

    man_name = change["fullDocument"]["managerName"]



    #image_path = "C:\\Users\\Administrator\\Desktop\\sitaralogo.PNG"
    image_path = "C:\\Users\\Administrator\\Desktop\\SiTaRA_Banner.PNG"

    create_title_page_TaRA(man_name, projectname,custom_title_path_TaRA,image_path)

    #path = 'C:\\Users\\Administrator\\Desktop\\PDF_FILES\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf'

    path = 'C:\\Users\\Administrator\\Documents\\Si-TaRA_outputfiles\\Si-TaRA_Threat_and_Risk_Analysis_AP.pdf'

    extra_pdf = PyPDF2.PdfReader(custom_title_path_TaRA)
    for page_num in range(len(extra_pdf.pages)):
        page = extra_pdf.pages[page_num]
        pdf_writer.add_page(page)

    
    main_pdf = PyPDF2.PdfReader(path)
    for page_num in range(len(main_pdf.pages)):
        page = main_pdf.pages[page_num]
        pdf_writer.add_page(page)
    


    
    main_pdf1 = PyPDF2.PdfReader(attacktrees)
    for page_num in range(len(main_pdf1.pages)):
        page = main_pdf1.pages[page_num]
        pdf_writer.add_page(page)


    
    main_pdf2 = PyPDF2.PdfReader(traceabilitypath)
    for page_num in range(len(main_pdf2.pages)):
        page = main_pdf2.pages[page_num]
        pdf_writer.add_page(page)

    
    # #pdf_path=os.path.join(extra_pdf)
    # for filename in pdf_filenames:
    #     pdf_name_with_extension = filename + '.pdf'  # Append '.pdf' to the filename
    #     pdf_path = os.path.join(folder_path, pdf_name_with_extension)
    #     if os.path.exists(pdf_path):
    #         pdf = PyPDF2.PdfReader(pdf_path)
    #         for page_num in range(len(pdf.pages)):
    #             page = pdf.pages[page_num]
    #             pdf_writer.add_page(page)
    #     else:
    #         # print(f"PDF '{pdf_name_with_extension}' does not exist in the folder.")
    #         continue

    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

    # print("output_path is", output_path)
    return output_path











#####################################################################################################################



def get_tara_results(query_string):
    """
    Execute the TARA query and store the results in a list.
    Create a data frame from the results list and pass it to the update_tra_and_rra_template function.
    """
    global security_query

    security_query = query_string

    MySQL_DATABASE = DataBase("tara_new")
    # print(query_string)


    # Execute the TARA query and store the results in a list
    tara_results = MySQL_DATABASE.execute(query_string)
    # Create a data frame from the results list
    #print("SQL DATA IS ",tara_results)
    df = pd.DataFrame(tara_results)
    df.fillna('', inplace=True)
    df.columns = columns_names
    df.columns = df.columns.astype(str)
    #present_column_name = column_map.get(sheet_names)
    # print(df)

    # for counter in column_map:
        
   # print("Value in Column map is ",column_map)
    #new_df=remove_empty_rows_for_dict(df,column_map[])
        # counter = counter+1
        

    
    #print("column names are",df)
    #df_new = df.loc[df.notnull()]
    #for key in column_map.keys():
     #   value = column_map[key]
      #  new_df=remove_empty_rows_for_dict(df,value)
    #print("columns values are ", column_map.values())
    
        
    for key, value in column_map.items():
        dict_values = value
        matching_columns = [col for col in df.columns if any(col in d_val for d_val in dict_values)]
        #print("Matching columns are: {}",matching_columns)
        #print("SSSSSSSSSS")        
        remove_empty_rows_for_dict(df,matching_columns)
        
    
    
def remove_empty_rows_for_dict(df, column_dict):
    #print('inside remove empty rows dict', df)
    #print("df before is {}",df)
    new_df = df.loc[:, column_dict]
    #print(new_df)
    mask =  (new_df != '')

    #print("df after is {}",new_df)

    new_df = new_df[mask.all(axis=1)]
    # print("after remove emopy rows",new_df)
    #print("new_df, columndict",new_df, column_dict)
    
    
    #Append df to excel 
    update_tra_and_rra_template(new_df, column_dict)
    #print(new_df)

    #create_pdf_from_df(new_df, 'C:\\Users\\Administrator\\Desktop\\df_to_table.pdf')

    #table_data = df_to_table(new_df)
 
# Create PDF with the table data
    #create_pdf(table_data, 'C:\\Users\\Administrator\\Desktop\\df_to_table.pdf')

    return new_df



#################################################################################################################################


# # Function to convert DataFrame to a table
# def df_to_table11(df):
#     table_data = [df.columns.tolist()] + df.values.tolist()

#     # Calculate column widths based on the length of the content
#     col_widths = [max(len(str(row[i])) for row in table_data) * 15 for i in range(len(table_data[0]))]

#     table = Table(table_data,colWidths=col_widths)

#     # Set table style
#     style = TableStyle([
#         ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Header row background color
#         ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Header text color
#         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center align all cells
#        ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Table gridlines
#         ('BACKGROUND', (0, 1), (-1, -1), colors.beige),  # Alternate row background color
#     ])

#     table.setStyle(style)
#     return table

# # Function to create PDF with DataFrame as a table
# def create_pdf_from_df11(df, output_pdf_path):
#     doc = SimpleDocTemplate(output_pdf_path, pagesize=letter)
#     elements = []

#     # Convert DataFrame to a table
#     table = df_to_table(df)
#     elements.append(table)

#     # Build PDF document
#     doc.build(elements)



# # Function to convert DataFrame to table
# def df_to_table(df):
#     table_data = [list(df.columns)] + df.values.tolist()
#     return table_data
 
# # Function to create a PDF with a table
# def create_pdf(table_data, pdf_filename):
#     doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
#     table_style = TableStyle([
#         ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
#         ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
#         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#         ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
#         ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
#         ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
#         ('GRID', (0, 0), (-1, -1), 1, colors.black)
#     ])
 
#     table = Table(table_data, repeatRows=1, style=table_style, colWidths=[doc.width / len(table_data[0]) * 1.1 for _ in table_data[0]])
#     table.hAlign = 'CENTER'
#     content = [table]
#     doc.build(content, canvasmaker=NumberedCanvas)
 
# class NumberedCanvas(canvas.Canvas):
#     def __init__(self, *args, **kwargs):
#         canvas.Canvas.__init__(self, *args, **kwargs)
#         self._saved_page_states = []
 
#     def showPage(self):
#         self._saved_page_states.append(dict(self.__dict__))
#         self._startPage()
 
#     def save(self):
#         """add page info to each page (page x of y)"""
#         num_pages = len(self._saved_page_states)
#         for state in self._saved_page_states:
#             self.__dict__.update(state)
#             self.draw_page_number(num_pages)
#             canvas.Canvas.showPage(self)
#         canvas.Canvas.save(self)
 
#     def draw_page_number(self, page_count):
#         self.setFont("Helvetica", 9)
#         self.drawRightString(200 * 2.54, 2.5 * 2.54, "Page %d of %d" % (self._pageNumber, page_count))






#################################################################################################################################







def update_tra_and_rra_template(data_frame,column_dict):
    """
    Load the TRA and RRA template workbook.
    Iterate through the sheets in the workbook and update the data in each sheet using the data in the data frame.
    Save the workbook.
    """
    
    global tara_file_name
    global starting_sheets
    global counter_i

    tmp_sheet_names = ['Assumptions', 'MUCs', 'DSsConsequences', 'SecGoals', 'ThreatsDSs', '','']

    if methodology == "Attack Potential":
        tmp_sheet_names[5] = 'SecurityNeeds_AP'
        tmp_sheet_names[6] = 'ThreatEvaluation_AP'
    else:
        tmp_sheet_names[5] = 'SecurityNeeds_LE'
        tmp_sheet_names[6] = 'ThreatEvaluation_LE'

    #liklidf= create_attack_vectors('select Rating from ble GROUP BY Name, Security_Properties')
    #print("likedf is", liklidf)



    if counter_i < len(tmp_sheet_names):
        if (tmp_sheet_names[counter_i] != "ThreatsDSs")  and (tmp_sheet_names[counter_i] != "SecurityNeeds_AP") and (tmp_sheet_names[counter_i] != "ThreatEvaluation_LE") and (tmp_sheet_names[counter_i] != "ThreatEvaluation_AP"):
            # print("counter value outside ds is", counter_i)
            # print(data_frame) 
            if (tmp_sheet_names[counter_i] == "DSsConsequences"):
                # print("inside if",tmp_sheet_names[counter_i])
                # print(data_frame[['e']])
                df_no_e_duplicates = data_frame.drop_duplicates()
                
                # print("\nDataFrame after eliminating duplicates:")
                # print(df_no_e_duplicates)

                append_df_to_excel(tara_file_name,df_no_e_duplicates,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=1,index=0,header=None)
                
            else:
                # print("inside else",tmp_sheet_names[counter_i])
                append_df_to_excel(tara_file_name,data_frame,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=1,index=0,header=None)
            # print("Sheet Name is ",tmp_sheet_names[counter_i])
            counter_i = counter_i+1
            
            
               
        elif tmp_sheet_names[counter_i] == "ThreatsDSs":
            # print("INSIDE DSS")
            #print("inside ", remove_empty_rows_for_dict(data_frame, column_dict))
            # remove_empty_rows_for_dict(data_frame, column_dict[0])
            # remove_empty_rows_for_dict(data_frame, column_dict[1])
            # remove_empty_rows_for_dict(data_frame, column_dict[2])
            #new_col =old_df.loc[:, 'l']
            df1=data_frame[['l']] 
            df2 = data_frame[['m']]
            df3 = data_frame[['n']]

            #print("Damage Scenario testing:",df3)
            #dftest=df3.drop_duplicates(keep=False,inplace=True)
            #print("Damage Scenario testing 2:",df3)


            # print("sheet names is ", tmp_sheet_names[counter_i])
            # #print("new", remove_empty_rows_for_dict(data_frame, column_dict[0]))
            append_df_to_excel(tara_file_name,df1,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=1,index=0,header=None)
            # #new_col =old_df.loc[:, 'm']
            #print("DAtaframe l is ", df1)
            # #print("new", remove_empty_rows_for_dict(data_frame, column_dict[1]))
            append_df_to_excel(tara_file_name,df2,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=2,index=0,header=None)
            # #new_col = old_df.loc[:, 'n']
            #print("Dataframe m is", df2)
            # #print("new", remove_empty_rows_for_dict(data_frame, column_dict[2]))

            append_df_to_excel(tara_file_name,df3,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=4,index=0,header=None)          
            # counter_i = 0
            #print("Dataframe n is", df3)
            # print("current counter value in dss ", counter_i)
            counter_i = counter_i + 1
            # print("counter value after in dss", counter_i)
    
            
        elif tmp_sheet_names[counter_i] == "SecurityNeeds_AP":
            
            # print("Inside SecurityNeeds_LE check")
            # print("Sheet Name is ",tmp_sheet_names[counter_i])
            
            df1=data_frame[['s']]
            df2=data_frame[['t']]
            # print(df1)
            #counter_i = counter_i + 1
            
            
            append_df_to_excel(tara_file_name,df1,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=1,index=0,header=None)
            append_df_to_excel(tara_file_name,df2,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=2,index=0,header=None)
            counter_i = counter_i+1

        elif tmp_sheet_names[6] == "ThreatEvaluation_LE":
            
            var1 = f"select Rating from `{interfaces_list_AP[0].strip().lower()}` GROUP BY Name, Security_Properties"
            
            for i in range(0,len(interfaces_list_AP)):
                if i==0:
                    continue
                else:
                    var1 = var1+f" union all select Rating from `{interfaces_list_AP[i].strip().lower()}` GROUP BY Name, Security_Properties"
            
            #liklidf= create_attack_vectors(f'select Rating from {interfaces_list_AP} GROUP BY Name, Security_Properties')
            # print(var1)
            liklidf= create_attack_vectors(var1)


            # print("Inside ThreatEvaluation_LE check")
            # print("Sheet Name is ",tmp_sheet_names[counter_i])
            
            #df1=data_frame[['u']]
            
            # print(liklidf)
            
            append_df_to_excel(tara_file_name,liklidf[['a']],sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=2,index=0,header=None)
            #append_df_to_excel(tara_file_name,df2,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=2,index=0,header=None)
            # counter_i = counter_i+1
            counter_i = counter_i + 1

        elif tmp_sheet_names[6] == "ThreatEvaluation_AP":
            #select Time, Expertise,Knowledge,Access,Equipment,Threat from `app service` where SUM IN(select MIN(SUM) from `app service` group by Name,Security_Properties) group by Name,Security_Properties;            
            var1 = f"select Time,Expertise,Knowledge,Access,Equipment from `{interfaces_list_AP[0].strip().lower()}` where SUM IN(select distinct MIN(SUM) from `{interfaces_list_AP[0].strip().lower()}` group by Name,Security_Properties) group by Name,Security_Properties"
            
            for i in range(0,len(interfaces_list_AP)):
                if i==0:
                    continue
                else:
                    var1 = var1+f" union all select Time,Expertise,Knowledge,Access,Equipment from `{interfaces_list_AP[i].strip().lower()}` where SUM IN(select distinct MIN(SUM) from `{interfaces_list_AP[i].strip().lower()}` group by Name,Security_Properties) group by Name,Security_Properties"
            
            #liklidf= create_attack_vectors(f'select Rating from {interfaces_list_AP} GROUP BY Name, Security_Properties')
            # print(var1)
            liklidf= create_attack_vectors_AP(var1)

            # print(liklidf)

            # print("Inside ThreatEvaluation_AP check")
            # print("Sheet Name is ",tmp_sheet_names[counter_i])
            
            df1=liklidf[['a']]
            df2=liklidf[['b']]
            df3=liklidf[['c']]
            df4=liklidf[['d']]
            df5=liklidf[['e']]
            
            # print(liklidf)
            
            append_df_to_excel(tara_file_name,df1,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_AP[0],startcol=2,index=0,header=None)
            append_df_to_excel(tara_file_name,df2,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_AP[0],startcol=3,index=0,header=None)
            append_df_to_excel(tara_file_name,df3,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_AP[0],startcol=4,index=0,header=None)
            append_df_to_excel(tara_file_name,df4,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_AP[0],startcol=5,index=0,header=None)
            append_df_to_excel(tara_file_name,df5,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_AP[0],startcol=6,index=0,header=None)
            
            
            
            
            #append_df_to_excel(tara_file_name,df2,sheet_name=tmp_sheet_names[counter_i],startrow=start_row_list[counter_i],startcol=2,index=0,header=None)
            # counter_i = counter_i+1
            counter_i = counter_i + 1

        
def update_workbook_with_macros(file_name):
    wb=load_workbook(file_name,keep_vba=True)
    wb.save(file_name)
    wb.close()
    


def change_stream_thread():
    global change
    print("inside change stream")
    change_stream = client.master_tara.Interfaces.watch(
        [
            {
                "$match": {
                    "operationType": {"$in": ["insert"]}
                }
            }
        ]
    )
    while True:
        for change in change_stream:
            global email
            global session_id
            session_id = change["fullDocument"]["sessionId"]
            current_status = change["fullDocument"]["status"]
            email = change["fullDocument"]["email"]
            print(change_stream.try_next())
            # print("current_status in change_stream", current_status)
            logging.debug(f'Change detected with session id: {session_id}, mail is: {email} and status: {current_status}')
            #update_tara_status(session_id, current_status)
            with open("C:\\Users\\Administrator\\Desktop\\Python-Codes\\session_id.txt", "w") as file:
                file.write(session_id)
            get_new_record(session_id)

# start the change stream thread
thread = threading.Thread(target=change_stream_thread)
thread.start()

 
# stop the change stream thread when the code exits
thread.join()






